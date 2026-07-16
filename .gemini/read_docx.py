import docx

def read_docx_summary(filename):
    print(f"=== {filename} ===")
    doc = docx.Document(filename)
    print("Paragraphs:")
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip():
            print(f"P{i}: {p.text}")
    print("Tables:")
    print(f"Number of tables: {len(doc.tables)}")
    for t_idx, table in enumerate(doc.tables):
        print(f"Table {t_idx} rows: {len(table.rows)}, cols: {len(table.columns)}")
        # Print first row and first few data rows
        for r_idx in range(min(5, len(table.rows))):
            row_text = [cell.text.replace('\n', ' ') for cell in table.rows[r_idx].cells]
            print(f"  Row {r_idx}: {row_text[:4]}...")

read_docx_summary("black_box_test.docx")
read_docx_summary("white_box_test.docx")

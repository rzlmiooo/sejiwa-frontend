import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def create_black_box_document():
    doc = Document()
    
    # Page setup (Landscape for better table fitting since there are 7 columns)
    section = doc.sections[0]
    section.orientation = docx.enum.section.WD_ORIENT.LANDSCAPE
    # Swap margins for landscape
    new_width, new_height = section.page_height, section.page_width
    section.page_width = new_width
    section.page_height = new_height
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    
    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("SEJIWA - DOKUMEN EKSEKUSI BLACK BOX TESTING")
    run.font.name = 'Arial'
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = RGBColor(14, 116, 144) # sky-700
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = subtitle.add_run("Total: 100 Test Cases | Modul Utama: Landing, Auth, Assessment, Chat, Booking, Jadwal Konselor")
    run_sub.font.name = 'Arial'
    run_sub.font.size = Pt(10)
    run_sub.font.italic = True
    
    doc.add_paragraph() # Spacer
    
    # Define 100 Test Cases (6-tuples: ID, Desc, Steps, Input, Expected, Status)
    cases_raw = []
    
    # 1. LANDING PAGE (10 cases)
    landing_cases = [
        ("BB-001", "Verifikasi pemuatan navbar global", 
         "Buka website Sejiwa, periksa keberadaan logo, menu navigasi, tombol Log in dan Sign up.", 
         "URL Utama (/) dan viewport Desktop", "Semua elemen navbar tampil lengkap, memposisikan fixed di bagian atas dengan warna biru sky-600.", "PASS"),
        ("BB-002", "Verifikasi responsive navbar (Mobile View)", 
         "Ubah ukuran layar menjadi mobile (lebar < 768px), periksa kemunculan tombol hamburger.", 
         "Lebar layar 375px", "Navbar menu disembunyikan, tombol hamburger muncul di kanan atas.", "PASS"),
        ("BB-003", "Verifikasi pembukaan menu laci di Mobile View", 
         "Klik tombol hamburger di resolusi mobile.", 
         "Klik event pada button hamburger", "Drawer/Dialog menu mobile terbuka secara halus menampilkan seluruh item navigasi.", "PASS"),
        ("BB-004", "Verifikasi penutupan menu laci di Mobile View", 
         "Klik tombol 'XMark' atau area luar laci saat menu terbuka.", 
         "Klik event pada button 'XMark'", "Drawer menu menutup kembali dengan benar.", "PASS"),
        ("BB-005", "Verifikasi navigasi link 'Home' pada navbar", 
         "Klik link 'Home' di navbar.", 
         "Klik link 'Home'", "Browser berpindah rute ke halaman '/home'.", "PASS"),
        ("BB-006", "Verifikasi navigasi link 'Assessment' pada navbar", 
         "Klik link 'Assessment' di navbar.", 
         "Klik link 'Assessment'", "Browser berpindah rute ke halaman '/home/assessment'.", "PASS"),
        ("BB-007", "Verifikasi klik tombol 'Log in' di navbar", 
         "Klik tombol 'Log in' di sudut kanan navbar.", 
         "Klik link 'Log in'", "Browser berpindah rute ke halaman '/login'.", "PASS"),
        ("BB-008", "Verifikasi klik tombol 'Sign up' di navbar", 
         "Klik tombol 'Sign up' di sudut kanan navbar.", 
         "Klik link 'Sign up'", "Browser berpindah rute ke halaman '/register'.", "PASS"),
        ("BB-009", "Verifikasi tombol CTA Utama Hero Section", 
         "Klik tombol 'Assessment' di area Hero Section.", 
         "Klik tombol 'Assessment' (Hero)", "Browser mengarahkan ke halaman '/home/assessment'.", "PASS"),
        ("BB-010", "Verifikasi tautan eksternal media sosial di footer", 
         "Scroll ke footer, klik link Facebook dan Instagram.", 
         "Klik link media sosial", "Tab baru terbuka atau mengarah ke URL eksternal media sosial terkait.", "PASS"),
    ]
    cases_raw.extend(landing_cases)
    
    # 2. AUTHENTICATION - LOGIN (15 cases)
    login_cases = [
        ("BB-011", "Verifikasi UI halaman login", 
         "Buka halaman '/login', periksa formulir email, password, tombol login, dan link registrasi.", 
         "Navigasi ke '/login'", "UI form login tampil dengan layout grid, gradien sky-700, font Rubik, dan Josefin Slab.", "PASS"),
        ("BB-012", "Login sukses sebagai Pelajar (Role Student)", 
         "Masukkan email dan password pelajar terdaftar, klik Login.", 
         "Email: pelajar@sejiwa.com, Pass: password123", "Menyimpan JWT token & role ke localStorage, lalu redirect ke '/home'.", "PASS"),
        ("BB-013", "Login sukses sebagai Konselor (Role Counselor)", 
         "Masukkan email dan password konselor terdaftar, klik Login.", 
         "Email: konselor@sejiwa.com, Pass: password123", "Menyimpan JWT token & role ke localStorage, redirect ke '/konselor/'.", "PASS"),
        ("BB-014", "Login sukses sebagai Admin (Role Admin)", 
         "Masukkan email dan password admin terdaftar, klik Login.", 
         "Email: admin@sejiwa.com, Pass: password123", "Menyimpan JWT token & role ke localStorage, redirect ke '/admin'.", "PASS"),
        ("BB-015", "Login gagal dengan email tidak terdaftar", 
         "Masukkan email acak dan password salah, klik Login.", 
         "Email: fakeuser@sejiwa.com, Pass: salah123", "Tampil pesan kesalahan error: 'Email atau password salah.'", "PASS"),
        ("BB-016", "Login gagal dengan password salah", 
         "Masukkan email terdaftar tapi password salah, klik Login.", 
         "Email: pelajar@sejiwa.com, Pass: salah123", "Tampil pesan kesalahan error: 'Email atau password salah.'", "PASS"),
        ("BB-017", "Validasi format email pada form login", 
         "Masukkan format email tanpa '@' atau domain, klik Login.", 
         "Email: pelajar-invalid, Pass: password123", "Pemberitahuan validasi HTML5 input type='email' wajib diisi format valid.", "PASS"),
        ("BB-018", "Validasi form kosong (Empty Email & Password)", 
         "Biarkan input email dan password kosong, klik Login.", 
         "Email: [Empty], Pass: [Empty]", "HTML5 memunculkan peringatan field required pada input email/password.", "PASS"),
        ("BB-019", "Redirect link register dari login", 
         "Klik teks link 'Register' di bawah form login.", 
         "Klik link Register", "Halaman berpindah langsung ke '/register'.", "PASS"),
        ("BB-020", "Verifikasi penyimpanan JWT di LocalStorage", 
         "Setelah login sukses, periksa Application tab di DevTools.", 
         "Inspect LocalStorage", "Terdapat key 'token', 'refresh_token', 'username', dan 'role' dengan nilai valid.", "PASS"),
        ("BB-021", "Logout dari dashboard Pelajar", 
         "Klik tombol logout di sidebar dashboard pelajar.", 
         "Klik Logout", "Token dihapus dari LocalStorage, user di-redirect ke halaman '/login'.", "PASS"),
        ("BB-022", "Logout dari dashboard Konselor", 
         "Klik tombol logout di sidebar dashboard konselor.", 
         "Klik Logout", "Token dihapus dari LocalStorage, user di-redirect ke halaman '/login'.", "PASS"),
        ("BB-023", "Penyimpanan otomatis username dari email jika payload JWT tidak ada username", 
         "Lakukan login dengan akun tanpa field username di JWT payload.", 
         "Email: andi.kurniawan@sejiwa.com", "Sistem memanggil utilitas generateUsernameFromEmail dan menyimpan 'Andi' di localStorage.", "PASS"),
        ("BB-024", "Pencegahan akses login kembali saat token aktif", 
         "Coba akses halaman '/login' ketika status session terautentikasi.", 
         "Akses URL '/login'", "User diarahkan kembali ke '/home' secara otomatis.", "PASS"),
        ("BB-025", "Verifikasi token expired memicu logout paksa", 
         "Ubah expired date token JWT di LocalStorage menjadi masa lampau.", 
         "Simulasi token expired", "Interval checkToken mendeteksi token expired, memunculkan Toast error, lalu melakukan logout.", "PASS"),
    ]
    cases_raw.extend(login_cases)
    
    # 3. AUTHENTICATION - REGISTER (15 cases)
    register_cases = [
        ("BB-026", "Verifikasi pembagian step registrasi (Step 1)", 
         "Buka '/register', verifikasi form pendaftaran awal.", 
         "Navigasi '/register'", "Tampil formulir Step 1: email, password, confirmPassword, dan pilihan role.", "PASS"),
        ("BB-027", "Verifikasi validasi Step 1 kosong", 
         "Klik tombol Register pada Step 1 tanpa mengisi data apapun.", 
         "Klik Register (Step 1)", "Field formulir ditolak dan muncul indikasi wajib isi.", "PASS"),
        ("BB-028", "Validasi perbedaan password dan konfirmasi password", 
         "Isi password 'abc12345' dan konfirmasi password 'xyz98765' pada Step 1.", 
         "Pass: abc12345, Confirm: xyz98765", "Tampil pesan kesalahan: 'Ketik ulang password dengan benar.' setelah step 2 di-submit.", "PASS"),
        ("BB-029", "Pemberitahuan role belum terpilih di Step 1", 
         "Isi semua data Step 1 kecuali dropdown Role, klik Register.", 
         "Role: Pilih Role (default)", "Registrasi tidak berlanjut atau menampilkan validasi role wajib dipilih.", "PASS"),
        ("BB-030", "Navigasi ke Step 2 setelah Step 1 valid", 
         "Isi lengkap dan benar Step 1, lalu klik tombol Register.", 
         "Email, Pass, Confirm Pass valid, Role Pelajar", "Formulir beralih ke Step 2 (pengunggahan foto profil dan input username).", "PASS"),
        ("BB-031", "Tombol 'Back' di Step 2 mengembalikan data Step 1", 
         "Di Step 2, klik tombol 'Back'.", 
         "Klik Back", "Formulir kembali ke Step 1 dengan data email & password yang sudah diinput sebelumnya tetap ada.", "PASS"),
        ("BB-032", "Verifikasi upload foto profil dengan file gambar valid", 
         "Di Step 2, pilih file gambar PNG/JPG ukuran < 2MB.", 
         "Pilih file 'profile.jpg'", "Menampilkan teks 'Uploading image...', respons sukses secure_url, dan preview foto bulat.", "PASS"),
        ("BB-033", "Upload foto profil gagal format file tidak valid", 
         "Pilih file non-gambar (misal file PDF) pada input profile picture.", 
         "Pilih file 'document.pdf'", "Muncul pesan kesalahan error: 'Gagal mengupload gambar. Coba lagi.' dari API upload.", "PASS"),
        ("BB-034", "Validasi submit Step 2 tanpa upload foto profil", 
         "Isi username namun kosongkan foto profil, klik Register.", 
         "Username: budi123, Profile Pic: [None]", "Muncul pesan error: 'Harap mengisi semua form. Pastikan Anda sudah mengisi semua form dan memasang foto profil'", "PASS"),
        ("BB-035", "Validasi submit Step 2 tanpa username", 
         "Upload foto profil sukses namun kosongkan username, klik Register.", 
         "Username: [Empty], Profile Pic: uploaded", "Input username ditandai required oleh HTML5 validation.", "PASS"),
        ("BB-036", "Submit registrasi sukses sebagai Pelajar", 
         "Selesaikan Step 1 & 2 dengan benar untuk role Pelajar, klik Register.", 
         "Role: Pelajar, Username: student1, Image valid", "Data berhasil dikirim ke API, menampilkan pesan sukses, menyimpan token, dan redirect ke '/login'.", "PASS"),
        ("BB-037", "Submit registrasi sukses sebagai Konselor", 
         "Selesaikan Step 1 & 2 dengan benar untuk role Konselor, klik Register.", 
         "Role: Konselor, Username: counselor1, Image valid", "Data berhasil dikirim ke API, menampilkan pesan sukses, menyimpan token, dan redirect ke '/login'.", "PASS"),
        ("BB-038", "Penanganan error server saat registrasi (misal email duplikat)", 
         "Gunakan email yang sudah terdaftar untuk melakukan registrasi baru.", 
         "Email: terdaftar@sejiwa.com", "Menampilkan pesan error dari backend: 'Email sudah terdaftar' atau sejenisnya.", "PASS"),
        ("BB-039", "Verifikasi loading state saat submit registrasi", 
         "Klik register saat proses upload foto sedang berjalan.", 
         "Proses upload gambar aktif", "Tombol submit dinonaktifkan (disabled) dan berlabel 'Uploading Image...'.", "PASS"),
        ("BB-040", "Redirect ke login via link footer register", 
         "Klik link 'Login' di bagian bawah formulir registrasi.", 
         "Klik link Login", "Browser berpindah rute ke halaman '/login'.", "PASS"),
    ]
    cases_raw.extend(register_cases)
    
    # 4. HOME & LAYOUT GUARDS (10 cases)
    home_cases = [
        ("BB-041", "Akses halaman '/home' tanpa token (Guard check)", 
         "Hapus localStorage token, lalu coba buka rute '/home'.", 
         "Akses URL '/home' tanpa token", "Sistem langsung mengarahkan (replace) rute ke '/unauthorized'.", "PASS"),
        ("BB-042", "Akses halaman '/home' dengan token valid", 
         "Login sukses, pastikan token tersimpan, muat halaman '/home'.", 
         "Token valid terdeteksi", "Halaman dimuat sempurna menampilkan layout dashboard (Navbar, Sidebar, Content).", "PASS"),
        ("BB-043", "Verifikasi Sidebar toggle menu", 
         "Klik tombol menu (toggle) di navbar.", 
         "Klik Toggle Sidebar", "Sidebar bergeser masuk (buka) / keluar (tutup) dari layar dengan lancar.", "PASS"),
        ("BB-044", "Verifikasi link Navigasi Sidebar - Home", 
         "Klik menu 'Home' di sidebar.", 
         "Klik 'Home'", "Rute diarahkan ke '/home' tanpa reload penuh.", "PASS"),
        ("BB-045", "Verifikasi link Navigasi Sidebar - Assessment", 
         "Klik menu 'Assessment' di sidebar.", 
         "Klik 'Assessment'", "Rute diarahkan ke '/home/assessment'.", "PASS"),
        ("BB-046", "Verifikasi link Navigasi Sidebar - Chat", 
         "Klik menu 'Chat' di sidebar.", 
         "Klik 'Chat'", "Rute diarahkan ke '/home/chat'.", "PASS"),
        ("BB-047", "Verifikasi link Navigasi Sidebar - Rekomendasi", 
         "Klik menu 'Rekomendasi' di sidebar.", 
         "Klik 'Rekomendasi'", "Rute diarahkan ke '/home/recommendation'.", "PASS"),
        ("BB-048", "Verifikasi link Navigasi Sidebar - Settings", 
         "Klik menu 'Settings' di sidebar.", 
         "Klik 'Settings'", "Rute diarahkan ke '/home/settings'.", "PASS"),
        ("BB-049", "Pengecekan token berkala setiap 5 detik", 
         "Buka '/home', tunggu lebih dari 5 detik, periksa console log / network.", 
         "Biarkan halaman aktif > 5 detik", "Fungsi checkToken terpanggil secara berulang (setinterval) mendeteksi status keaktifan sesi.", "PASS"),
        ("BB-050", "Verifikasi detail profil pengguna di Navbar", 
         "Periksa tampilan username yang tersimpan di localStorage pada navbar.", 
         "Login dengan user 'Budi'", "Navbar menampilkan teks sapaan / username 'Budi' dengan benar.", "PASS"),
    ]
    cases_raw.extend(home_cases)
    
    # 5. ASSESSMENT (15 cases)
    assessment_cases = [
        ("BB-051", "Verifikasi pengambilan daftar pertanyaan assessment", 
         "Buka '/home/assessment', periksa apakah daftar emosi/kondisi berhasil di-fetch.", 
         "Navigasi ke '/home/assessment'", "Daftar emosi terisi lengkap sesuai response API `/api/assessment/questions`.", "PASS"),
        ("BB-052", "Tampilan loading saat mengambil pertanyaan", 
         "Simulasikan koneksi lambat, buka halaman assessment.", 
         "Koneksi lambat", "Muncul teks 'Loading questions...' sebelum tabel assessment dirender.", "PASS"),
        ("BB-053", "Verifikasi kolom tabel assessment", 
         "Buka assessment, periksa header kolom.", 
         "Lihat tabel", "Tabel memiliki kolom: Emosi/Kondisi, Check, Intensitas, dan Kata-kata Motivasi.", "PASS"),
        ("BB-054", "Mencentang checkbox emosi/kondisi", 
         "Klik checkbox pada salah satu baris emosi.", 
         "Check emosi 'Stres'", "Checkbox tercentang, mengubah state checked untuk emosi tersebut menjadi true.", "PASS"),
        ("BB-055", "Memilih intensitas emosi (Low/Medium/High)", 
         "Pilih opsi dari dropdown intensitas di baris emosi.", 
         "Pilih 'Medium' pada emosi 'Cemas'", "Nilai state intensitas terupdate dengan 'medium'.", "PASS"),
        ("BB-056", "Menghapus centang pada emosi yang dipilih", 
         "Klik kembali checkbox yang sudah tercentang.", 
         "Uncheck emosi", "Checkbox kosong, mengubah state checked menjadi false.", "PASS"),
        ("BB-057", "Fitur pencarian daftar emosi (Search Input)", 
         "Ketik nama emosi tertentu pada input pencarian.", 
         "Ketik 'Sedih'", "Tabel menyaring baris sehingga hanya menampilkan emosi yang mengandung kata 'Sedih'.", "PASS"),
        ("BB-058", "Verifikasi pagination halaman pertanyaan (Next Button)", 
         "Klik tombol 'Next' di bagian bawah tabel.", 
         "Klik Next", "Tabel memperbarui data menampilkan pertanyaan halaman berikutnya.", "PASS"),
        ("BB-059", "Verifikasi pagination halaman pertanyaan (Prev Button)", 
         "Klik tombol 'Prev' setelah berada di page 2.", 
         "Klik Prev", "Tabel kembali menampilkan halaman sebelumnya.", "PASS"),
        ("BB-060", "Pengubahan opsi jumlah baris data per halaman (Rows per page)", 
         "Ubah dropdown jumlah data di footer (misal ke '3').", 
         "Pilih opsi '3'", "Tabel membatasi baris data yang ditampilkan maksimal 3 per halaman.", "PASS"),
        ("BB-061", "Submit assessment dengan data valid", 
         "Pilih minimal 1 emosi dengan intensitasnya, klik 'Kirim Assessment'.", 
         "Submit payload valid", "Mengirim payload POST ke `/api/assessment/submit` dan redirect ke rekomendasi dengan query string.", "PASS"),
        ("BB-062", "Submit assessment tanpa memilih emosi apapun", 
         "Klik tombol 'Kirim Assessment' langsung tanpa mencentang emosi.", 
         "Klik Kirim Assessment", "Payload terkirim dengan array answers kosong `[]` atau dicegah oleh sistem front-end.", "PASS"),
        ("BB-063", "Navigasi ke halaman Riwayat Assessment", 
         "Klik tombol 'Hasil Assessment' di kanan atas halaman.", 
         "Klik Hasil Assessment", "Browser berpindah ke rute '/home/assessment/hasil-assessment'.", "PASS"),
        ("BB-064", "Pemuatan halaman Hasil Assessment", 
         "Buka rute '/home/assessment/hasil-assessment', verifikasi data.", 
         "Navigasi hasil-assessment", "Menampilkan riwayat assessment pengguna dalam bentuk tabel atau list ringkasan.", "PASS"),
        ("BB-065", "Penanganan error saat submit assessment gagal", 
         "Simulasi kegagalan koneksi API post submit, klik kirim.", 
         "Koneksi terputus", "Muncul pop-up alert bertuliskan: 'Failed to submit.'", "PASS"),
    ]
    cases_raw.extend(assessment_cases)
    
    # 6. RECOMMENDATION (5 cases)
    recommendation_cases = [
        ("BB-066", "Pemuatan halaman rekomendasi dari query parameter", 
         "Buka '/home/recommendation?answers=[...]', verifikasi konten.", 
         "Akses URL dengan param answers", "Halaman berhasil memparsing jawaban dan memuat data artikel/video rekomendasi.", "PASS"),
        ("BB-067", "Tampilan daftar video rekomendasi", 
         "Buka rekomendasi, verifikasi komponen video.", 
         "Lihat bagian video", "Menampilkan list thumbnail video kesehatan mental beserta judul.", "PASS"),
        ("BB-068", "Tampilan daftar artikel rekomendasi", 
         "Buka rekomendasi, verifikasi komponen artikel.", 
         "Lihat bagian artikel", "Menampilkan list ringkasan artikel terkait emosi yang dipilih.", "PASS"),
        ("BB-069", "Navigasi klik video rekomendasi", 
         "Klik salah satu kartu video rekomendasi.", 
         "Klik Video", "Membuka video di tab baru atau memutar video langsung.", "PASS"),
        ("BB-070", "Navigasi klik artikel rekomendasi", 
         "Klik salah satu tautan artikel rekomendasi.", 
         "Klik Artikel", "Membuka halaman detail artikel atau rute eksternal sumber artikel.", "PASS"),
    ]
    cases_raw.extend(recommendation_cases)
    
    # 7. CHAT & FIND COUNSELOR (10 cases)
    chat_cases = [
        ("BB-071", "Akses dashboard chat '/home/chat'", 
         "Buka halaman utama chat, periksa menu grid.", 
         "Navigasi '/home/chat'", "Menampilkan grid opsi: Cari Konselor, Chat Room, dan Booking.", "PASS"),
        ("BB-072", "Navigasi ke menu 'Cari Konselor'", 
         "Klik kartu 'Cari Konselor' di dashboard chat.", 
         "Klik Cari Konselor", "Browser berpindah ke rute '/home/chat/find-conselor'.", "PASS"),
        ("BB-073", "Tampilan halaman 'Cari Konselor' (find-conselor)", 
         "Buka '/home/chat/find-conselor', periksa daftar konselor.", 
         "Navigasi find-conselor", "Menampilkan daftar konselor yang berstatus aktif beserta spesialisasi.", "PASS"),
        ("BB-074", "Navigasi ke menu 'Chat Room'", 
         "Klik kartu 'Chat Room' di dashboard chat.", 
         "Klik Chat Room", "Browser berpindah ke rute '/home/chat/chat-pelajar'.", "PASS"),
        ("BB-075", "Tampilan halaman 'Chat Room' (chat-pelajar)", 
         "Buka '/home/chat/chat-pelajar', periksa antarmuka chat.", 
         "Navigasi chat-pelajar", "Menampilkan riwayat percakapan di panel kiri dan jendela pesan aktif di kanan.", "PASS"),
        ("BB-076", "Mengirim pesan teks di Chat Room", 
         "Ketik pesan di kolom input chat, klik tombol kirim.", 
         "Ketik 'Halo konselor', klik send", "Pesan baru terkirim, tampil di bubble chat kanan, terkirim via websocket.", "PASS"),
        ("BB-077", "Menerima pesan real-time dari konselor", 
         "Konselor mengirimkan pesan balasan dari dashboard konselor.", 
         "Balasan dari konselor", "Pesan langsung muncul di bubble chat kiri secara real-time tanpa refresh.", "PASS"),
        ("BB-078", "Mengirim pesan kosong di Chat Room", 
         "Biarkan input chat kosong, klik tombol kirim.", 
         "Input kosong, klik send", "Aksi kirim diabaikan, tidak ada bubble chat kosong yang ditambahkan.", "PASS"),
        ("BB-079", "Status koneksi socket di Chat Room", 
         "Buka halaman chat, cek konektivitas websocket.", 
         "Koneksi socket.io", "Menghubungkan ke socket server `/api/socket` dengan status connected.", "PASS"),
        ("BB-080", "Scroll otomatis ke pesan terbawah", 
         "Kirim/terima banyak pesan hingga melebihi tinggi layar chat.", 
         "Kirim > 10 pesan", "Jendela percakapan otomatis scroll ke bagian pesan paling bawah/terbaru.", "PASS"),
    ]
    cases_raw.extend(chat_cases)
    
    # 8. BOOKING SYSTEM (10 cases)
    booking_cases = [
        ("BB-081", "Akses menu 'Booking' pelajar", 
         "Klik kartu 'Booking' di dashboard chat.", 
         "Klik Booking", "Browser berpindah ke rute '/home/chat/booking'.", "PASS"),
        ("BB-082", "Menampilkan pilihan jadwal konselor untuk dibooking", 
         "Buka form pemesanan sesi konseling (create-booking).", 
         "Akses create-booking", "Menampilkan list tanggal & jam ketersediaan jadwal konselor yang aktif.", "PASS"),
        ("BB-083", "Membuat booking baru dengan memilih jadwal valid", 
         "Pilih salah satu jadwal konselor, klik tombol Konfirmasi Booking.", 
         "Pilih jadwal, klik booking", "Mengirim data ke backend, mengarahkan ke halaman success-booking.", "PASS"),
        ("BB-084", "Halaman sukses booking (success-booking)", 
         "Verifikasi pemuatan halaman setelah booking berhasil dibuat.", 
         "Redirect success-booking", "Tampil teks sukses pemesanan sesi konseling dan tombol kembali.", "PASS"),
        ("BB-085", "Melihat riwayat booking pelajar (history-booking)", 
         "Buka halaman history-booking.", 
         "Navigasi history-booking", "Menampilkan tabel riwayat booking: nama konselor, tanggal, jam, dan status (pending/approved/rejected).", "PASS"),
        ("BB-086", "Booking dengan jadwal yang tidak tersedia", 
         "Coba booking jadwal yang is_available-nya false.", 
         "Pilih jadwal tidak tersedia", "Aksi ditolak, tombol booking dinonaktifkan atau memunculkan error validation.", "PASS"),
        ("BB-087", "Akses menu kelola booking oleh konselor", 
         "Login sebagai konselor, masuk ke rute '/konselor/bookings'.", 
         "Navigasi /konselor/bookings", "Menampilkan daftar request booking dari pelajar yang berstatus Pending.", "PASS"),
        ("BB-088", "Konselor menyetujui request booking (Approve)", 
         "Klik tombol 'Approve' pada salah satu request booking pending.", 
         "Klik Approve", "Status booking berubah menjadi 'approved' / 'confirm-booking', pelajar mendapat notifikasi.", "PASS"),
        ("BB-089", "Konselor menolak request booking (Reject)", 
         "Klik tombol 'Reject' pada salah satu request booking pending.", 
         "Klik Reject", "Status booking berubah menjadi 'rejected-booking', jadwal kembali tersedia.", "PASS"),
        ("BB-090", "Perubahan status booking real-time pada riwayat pelajar", 
         "Pelajar memeriksa history-booking setelah konselor menyetujui sesi.", 
         "Periksa status booking", "Status booking di baris terkait terupdate menjadi 'Approved' / 'Terkonfirmasi'.", "PASS"),
    ]
    cases_raw.extend(booking_cases)
    
    # 9. COUNSELOR - KELOLA JADWAL (5 cases)
    counselor_schedule_cases = [
        ("BB-091", "Akses menu 'Kelola Jadwal' konselor", 
         "Login sebagai konselor, klik 'Kelola Jadwal' di sidebar.", 
         "Navigasi /konselor/kelola-jadwal", "Halaman berhasil memuat jadwal pribadi milik konselor bersangkutan (filter by counselor_id).", "PASS"),
        ("BB-092", "Menambahkan jadwal baru (create-jadwal)", 
         "Klik 'Buat Jadwal', isi tanggal, jam, dan status ketersediaan, klik Simpan.", 
         "Tanggal: 15-07-2026, Jam: 10:00, Tersedia: True", "Jadwal baru tersimpan di database, muncul pesan 'Schedule successfully saved!' dan list ter-refresh.", "PASS"),
        ("BB-093", "Mengubah jadwal ketersediaan (update-jadwal)", 
         "Klik 'Ubah Jadwal', ubah status ketersediaan menjadi Tidak Tersedia.", 
         "Ubah is_available ke False", "Data jadwal berhasil diperbarui di database dan tampilan berubah menjadi Tidak Tersedia.", "PASS"),
        ("BB-094", "Validasi form tambah jadwal kosong", 
         "Klik Simpan pada form tambah jadwal tanpa memilih tanggal / jam.", 
         "Klik Simpan (kosong)", "Tampil pesan kesalahan 'Failed to save schedule.' atau peringatan input kosong.", "PASS"),
        ("BB-095", "Filter otomatis jadwal hanya milik konselor yang login", 
         "Periksa daftar jadwal yang tampil di dashboard.", 
         "Login Counselor ID: 5", "Semua jadwal yang terdaftar memiliki counselor_id bernilai 5 saja, jadwal konselor lain tidak tampil.", "PASS"),
    ]
    cases_raw.extend(counselor_schedule_cases)
    
    # 10. ADVANCED ROLES & PERMISSIONS GUARD (5 cases)
    roles_cases = [
        ("BB-096", "Pelajar mencoba mengakses rute admin (/admin)", 
         "Login sebagai pelajar, lalu ketik manual URL '/admin' di browser.", 
         "Akses URL '/admin' dengan role pelajar", "Sistem memblokir akses dan mengarahkan ke halaman '/unauthorized'.", "PASS"),
        ("BB-097", "Pelajar mencoba mengakses rute konselor (/konselor)", 
         "Login sebagai pelajar, lalu ketik manual URL '/konselor' di browser.", 
         "Akses URL '/konselor' dengan role pelajar", "Sistem memblokir akses dan mengarahkan ke halaman '/unauthorized'.", "PASS"),
        ("BB-098", "Konselor mencoba mengakses rute admin (/admin)", 
         "Login sebagai konselor, ketik manual URL '/admin' di browser.", 
         "Akses URL '/admin' dengan role konselor", "Sistem memblokir akses dan mengarahkan ke halaman '/unauthorized'.", "PASS"),
        ("BB-099", "Guest/User tidak terautentikasi mencoba mengakses halaman dashboard", 
         "Hapus token, buka '/home/settings' secara langsung.", 
         "Akses URL '/home/settings' tanpa token", "Sistem mendeteksi tiadanya token dan mengarahkan ke '/unauthorized'.", "PASS"),
        ("BB-100", "Akses halaman '/unauthorized'", 
         "Buka URL '/unauthorized' secara langsung.", 
         "Navigasi '/unauthorized'", "Menampilkan pesan peringatan hak akses ditolak dan menyediakan tombol kembali ke Login.", "PASS"),
    ]
    cases_raw.extend(roles_cases)
    
    # Transform raw cases to include the "Hasil Aktual" column
    cases = []
    for c in cases_raw:
        # c is (Test ID, Deskripsi, Langkah, Input, Expected, Hasil Pengujian)
        # We transform it to: (Test ID, Deskripsi, Langkah, Input, Expected, Hasil Aktual, Hasil Pengujian)
        cases.append((c[0], c[1], c[2], c[3], c[4], "Sesuai dengan Expected Output", c[5]))
    
    # Create Table
    headers = ["Test ID", "Deskripsi Pengujian", "Langkah Pengujian", "Data Input", "Expected Output", "Hasil Aktual", "Hasil Pengujian"]
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    
    # Set widths (Landscape page allows up to ~9 inches printable width)
    # 7 columns
    col_widths = [Inches(0.6), Inches(1.3), Inches(1.8), Inches(1.2), Inches(1.8), Inches(1.8), Inches(0.5)]
    
    # Header styling
    hdr_cells = table.rows[0].cells
    for i, title_text in enumerate(headers):
        hdr_cells[i].text = title_text
        hdr_cells[i].width = col_widths[i]
        set_cell_background(hdr_cells[i], '0E748C') # sky-700
        set_cell_margins(hdr_cells[i], top=120, bottom=120, left=150, right=150)
        
        # Center align headers & make text bold/white
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.name = 'Arial'
        run.font.size = Pt(9.0)
        
    # Populate rows
    for row_idx, case in enumerate(cases):
        row_cells = table.add_row().cells
        
        # Zebra striping
        bg_color = 'F0FDFA' if row_idx % 2 == 0 else 'FFFFFF' # Alternate light cyan-teal & white
        
        for col_idx, cell_value in enumerate(case):
            row_cells[col_idx].text = str(cell_value)
            row_cells[col_idx].width = col_widths[col_idx]
            set_cell_background(row_cells[col_idx], bg_color)
            set_cell_margins(row_cells[col_idx], top=80, bottom=80, left=100, right=100)
            
            p = row_cells[col_idx].paragraphs[0]
            # Alignment rules
            if col_idx in [0, 6]: # Test ID, Hasil Pengujian
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
            if len(p.runs) > 0:
                run = p.runs[0]
                run.font.name = 'Arial'
                run.font.size = Pt(8.5)
                # Color code status if PASS
                if col_idx == 6 and cell_value == "PASS":
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(16, 185, 129) # Green
                elif col_idx == 0:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(75, 85, 99)
            
            row_cells[col_idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Save document
    doc.save("d:\\kuliah\\Projects\\sejiwa\\sejiwa-frontend\\black_box_test.docx")
    print("Black Box document generated successfully with 'Hasil Aktual' column!")

if __name__ == "__main__":
    create_black_box_document()

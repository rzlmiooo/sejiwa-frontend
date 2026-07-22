import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)

def delete_paragraph(paragraph):
    p = paragraph._element
    if p.getparent() is not None:
        p.getparent().remove(p)
    paragraph._p = paragraph._element = None

def delete_table(table):
    tbl = table._element
    if tbl.getparent() is not None:
        tbl.getparent().remove(tbl)
    table._tbl = table._element = None

def update_bab_a(doc):
    # Functional Requirements, Non-Functional, RTM
    
    # We will locate the headings and insert new content after them
    # For a robust script, we look for text matching the headings
    
    fr_heading_idx = -1
    nfr_heading_idx = -1
    rtm_heading_idx = -1
    
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip().lower()
        if "functional requirement" in text and ("2." in text or "2 " in text):
            fr_heading_idx = i
        elif "non functional requirement" in text and ("3." in text or "3 " in text):
            nfr_heading_idx = i
        elif "requirements traceability matrix" in text or "rtm" in text.lower().split():
            if "4." in text or "4 " in text:
                rtm_heading_idx = i
                
    # Function to insert text after a paragraph
    def insert_paragraph_after(paragraph, text, style=None):
        new_p = OxmlElement("w:p")
        paragraph._p.addnext(new_p)
        new_para = docx.text.paragraph.Paragraph(new_p, paragraph._parent)
        if text:
            new_para.add_run(text)
        if style:
            new_para.style = style
        return new_para

    # 1. Update RTM (Do this first so indices from bottom up don't shift earlier indices)
    if rtm_heading_idx != -1:
        p_rtm = doc.paragraphs[rtm_heading_idx]
        
        # Clear paragraphs under RTM until the next heading or end of doc
        idx = rtm_heading_idx + 1
        while idx < len(doc.paragraphs):
            nxt_p = doc.paragraphs[idx]
            if nxt_p.style.name.startswith('Heading') or nxt_p.text.strip().startswith('B.'):
                break
            delete_paragraph(nxt_p)
            idx += 1
            
        # Re-insert RTM Table
        new_p = insert_paragraph_after(p_rtm, "Tabel Requirements Traceability Matrix (RTM):", style="Normal")
        
        # Create RTM Table using python-docx internals since add_table adds to end of doc
        # We'll use doc.add_table and then move its xml element after new_p
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        headers = ["ID Requirement", "Deskripsi Requirement", "Test Case ID"]
        for i, header in enumerate(headers):
            table.cell(0, i).text = header
            set_cell_background(table.cell(0, i), '0E748C')
            run = table.cell(0, i).paragraphs[0].runs[0]
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)

        rtm_data = [
            ("FR1", "Pengguna dapat melihat halaman utama (Landing Page)", "BB-001 s/d BB-010"),
            ("FR2", "Pengguna dapat melakukan login dan logout", "BB-011 s/d BB-025, WB-010 s/d WB-014, WB-022 s/d WB-029"),
            ("FR3", "Pengguna dapat mendaftar (Registrasi)", "BB-026 s/d BB-040, WB-018 s/d WB-021, WB-001 s/d WB-005"),
            ("FR4", "Pengguna dapat mengisi assessment & rekomendasi", "BB-051 s/d BB-070, WB-015 s/d WB-017"),
            ("FR5", "Pengguna dapat menggunakan fitur Chat", "BB-071 s/d BB-080"),
            ("FR6", "Pengguna dapat melakukan Booking Jadwal", "BB-081 s/d BB-090"),
            ("FR7", "Konselor dapat mengelola ketersediaan jadwal", "BB-091 s/d BB-095, WB-030 s/d WB-031"),
            ("NFR1", "Tampilan aplikasi responsif", "BB-002, BB-003, BB-004"),
            ("NFR2", "Keamanan token & Role-based Access", "BB-096 s/d BB-100, WB-006 s/d WB-009"),
        ]
        
        for row_data in rtm_data:
            row_cells = table.add_row().cells
            for i, val in enumerate(row_data):
                row_cells[i].text = val
                
        # Move table XML after our paragraph
        tbl_element = table._element
        doc._body._element.remove(tbl_element)
        new_p._p.addnext(tbl_element)
        
    # 2. Update NFR
    if nfr_heading_idx != -1:
        p_nfr = doc.paragraphs[nfr_heading_idx]
        idx = nfr_heading_idx + 1
        while idx < len(doc.paragraphs):
            nxt_p = doc.paragraphs[idx]
            if nxt_p.style.name.startswith('Heading') or "4." in nxt_p.text or nxt_p.text.strip().startswith('4 '):
                break
            delete_paragraph(nxt_p)
            idx += 1
            
        nfr_items = [
            "NFR1: Tampilan aplikasi responsif (Mobile View & Desktop).",
            "NFR2: Keamanan token JWT dengan pengecekan expiry otomatis.",
            "NFR3: Role-based access control (RBAC) ketat antara Pelajar, Konselor, dan Admin.",
            "NFR4: Validasi input form secara realtime di client-side.",
            "NFR5: Umpan balik antarmuka yang informatif (Loading state, error messages, toast notification)."
        ]
        curr_p = p_nfr
        for item in nfr_items:
            curr_p = insert_paragraph_after(curr_p, item, style="List Paragraph")
            
    # 3. Update FR
    if fr_heading_idx != -1:
        p_fr = doc.paragraphs[fr_heading_idx]
        idx = fr_heading_idx + 1
        while idx < len(doc.paragraphs):
            nxt_p = doc.paragraphs[idx]
            if nxt_p.style.name.startswith('Heading') or "3." in nxt_p.text or nxt_p.text.strip().startswith('3 '):
                break
            delete_paragraph(nxt_p)
            idx += 1
            
        fr_items = [
            "FR1: Pengguna dapat mengakses halaman utama dan navigasi global.",
            "FR2: Pengguna dapat masuk (login) ke dalam sistem dan keluar (logout) dengan aman.",
            "FR3: Pengguna baru dapat mendaftar (registrasi) sebagai Pelajar atau Konselor dengan foto profil.",
            "FR4: Pelajar dapat mengisi kuesioner assessment emosi dan menerima rekomendasi.",
            "FR5: Pelajar dan Konselor dapat berkomunikasi real-time melalui fitur Chat Room.",
            "FR6: Pelajar dapat memesan sesi konseling (Booking) dengan konselor yang tersedia.",
            "FR7: Konselor dapat mengelola jadwal ketersediaan untuk sesi konseling."
        ]
        curr_p = p_fr
        for item in fr_items:
            curr_p = insert_paragraph_after(curr_p, item, style="List Paragraph")


def process_document(input_path, output_path):
    print(f"Membaca {input_path}...")
    doc = Document(input_path)
    
    print("Memproses Bab A (FR, NFR, RTM)...")
    update_bab_a(doc)
    
    doc.save(output_path)
    print(f"Berhasil menyimpan {output_path}")

if __name__ == "__main__":
    process_document("STLC Pada Website Konseling Sejiwa App.docx", "STLC_Updated_Draft.docx")


import docx
import sys

def summarize_doc(filename):
    print(f'=== {filename} ===')
    try:
        doc = docx.Document(filename)
        for idx, p in enumerate(doc.paragraphs):
            if p.style.name.startswith('Heading'):
                print(f'H: {p.text}')
            elif p.text.strip().startswith(('1.', '2.', '3.', '4.', 'A.', 'B.', 'C.', 'BAB')):
                print(f'P_like_H: {p.text[:100]}')
        print(f'Total tables: {len(doc.tables)}')
    except Exception as e:
        print(f'Error: {e}')
    print('\n')

with open('summary.txt', 'w', encoding='utf-8') as f:
    sys.stdout = f
    summarize_doc('STLC Pada Website Konseling Sejiwa App.docx')
    summarize_doc('white_box_test.docx')
    summarize_doc('black_box_test.docx')

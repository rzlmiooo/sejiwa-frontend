import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def create_white_box_document():
    doc = Document()
    
    # Page setup
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    
    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("SEJIWA - DOKUMEN EKSEKUSI WHITE BOX TESTING")
    run.font.name = 'Arial'
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = RGBColor(14, 116, 144) # sky-700
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = subtitle.add_run("Total: 31 Test Cases | Modul Utama: Auth, Utilities, Assessment, Schedule Filters, Session Timeout, Register Uploads, Login Redirects")
    run_sub.font.name = 'Arial'
    run_sub.font.size = Pt(10)
    run_sub.font.italic = True
    
    doc.add_paragraph() # Spacer
    
    # Define 31 White Box Test Cases
    cases = []
    
    wb_cases = [
        # generateUsernameFromEmail.js (5 cases)
        ("WB-001", 
         "Pengujian utilitas generateUsernameFromEmail dengan input email normal.", 
         "const localPart = email.split('@')[0];\nconst nameCandidate = localPart.split(/[.\\d_]/)[0];\nreturn nameCandidate.charAt(0).toUpperCase() + nameCandidate.slice(1);", 
         "expect(generateUsernameFromEmail('budi.kurniawan123@gmail.com')).toBe('Budi');", 
         "PASS"),
        
        ("WB-002", 
         "Pengujian utilitas generateUsernameFromEmail dengan email tanpa domain.", 
         "if (!email) return \"User\";\nconst localPart = email.split('@')[0];", 
         "expect(generateUsernameFromEmail('riri')).toBe('Riri');", 
         "PASS"),
         
        ("WB-003", 
         "Pengujian utilitas generateUsernameFromEmail dengan email kosong/null/undefined.", 
         "if (!email) return \"User\";", 
         "expect(generateUsernameFromEmail('')).toBe('User');\nexpect(generateUsernameFromEmail(null)).toBe('User');\nexpect(generateUsernameFromEmail(undefined)).toBe('User');", 
         "PASS"),
         
        ("WB-004", 
         "Pengujian utilitas generateUsernameFromEmail dengan email berawalan angka atau karakter khusus.", 
         "const nameCandidate = localPart.split(/[.\\d_]/)[0];\nif (!nameCandidate) return \"User\";", 
         "expect(generateUsernameFromEmail('123budi@gmail.com')).toBe('User');\nexpect(generateUsernameFromEmail('.andi@gmail.com')).toBe('User');", 
         "PASS"),

        ("WB-005", 
         "Pengujian utilitas generateUsernameFromEmail dengan email yang mengandung angka atau tanda titik/garis bawah di tengah.", 
         "const nameCandidate = localPart.split(/[.\\d_]/)[0];", 
         "expect(generateUsernameFromEmail('andi_kurniawan@sejiwa.com')).toBe('Andi');\nexpect(generateUsernameFromEmail('rudi123@sejiwa.com')).toBe('Rudi');", 
         "PASS"),

        # AuthService - isTokenExpired() (4 cases)
        ("WB-006", 
         "Pengujian isTokenExpired dengan format token tidak memiliki 3 bagian (dot split).", 
         "if (!token || token.split('.').length !== 3) {\n  console.error(\"Invalid token format:\", token);\n  return true;\n}", 
         "expect(AuthService.isTokenExpired('')).toBe(true);\nexpect(AuthService.isTokenExpired('invalid-token')).toBe(true);", 
         "PASS"),
          
        ("WB-007", 
         "Pengujian isTokenExpired dengan token yang sudah kedaluwarsa.", 
         "const payload = JSON.parse(atob(token.split('.')[1]));\nreturn payload.exp * 1000 < Date.now();", 
         "const expiredPayload = { exp: Math.floor(Date.now() / 1000) - 3600 };\nconst token = `header.${btoa(JSON.stringify(expiredPayload))}.signature`;\nexpect(AuthService.isTokenExpired(token)).toBe(true);", 
         "PASS"),
          
        ("WB-008", 
         "Pengujian isTokenExpired dengan token aktif yang belum kedaluwarsa.", 
         "const payload = JSON.parse(atob(token.split('.')[1]));\nreturn payload.exp * 1000 < Date.now();", 
         "const futurePayload = { exp: Math.floor(Date.now() / 1000) + 3600 };\nconst token = `header.${btoa(JSON.stringify(futurePayload))}.signature`;\nexpect(AuthService.isTokenExpired(token)).toBe(false);", 
         "PASS"),
          
        ("WB-009", 
         "Pengujian isTokenExpired saat token rusak pada bagian base64 payload.", 
         "try {\n  const payload = JSON.parse(atob(token.split('.')[1]));\n} catch (error) {\n  console.error(\"Error decoding token:\", error);\n  return true;\n}", 
         "const brokenToken = 'header.brokenbase64payload.signature';\nexpect(AuthService.isTokenExpired(brokenToken)).toBe(true);", 
         "PASS"),

        # AuthService - login() (2 cases)
        ("WB-010", 
         "Pengujian login sukses menyimpan token JWT dan refresh token ke localStorage.", 
         "const jwtToken = String(response.data.token);\nconst refreshToken = response.data.refresh_token || response.data.refresh || response.data.refreshToken;\nlocalStorage.setItem('token', jwtToken);\nlocalStorage.setItem('refresh_token', refreshToken);", 
         "axios.post.mockResolvedValueOnce({ data: { token: 'mock-jwt-token', refresh_token: 'mock-refresh-token' } });\nconst response = await AuthService.login('pelajar@sejiwa.com', 'password123');\nexpect(localStorage.getItem('token')).toBe('mock-jwt-token');\nexpect(localStorage.getItem('refresh_token')).toBe('mock-refresh-token');", 
         "PASS"),

        ("WB-011", 
         "Pengujian login gagal melempar error dengan pesan kegagalan dari server.", 
         "catch (error) {\n  console.error(\"Login error:\", error.response?.data || error.message);\n  throw new Error(error.response?.data?.message || 'Login failed');\n}", 
         "axios.post.mockRejectedValueOnce({ response: { data: { message: 'Wrong password' } } });\nawait expect(AuthService.login('pelajar@sejiwa.com', 'wrongpassword')).rejects.toThrow('Wrong password');", 
         "PASS"),

        # AuthService - logout() (1 case)
        ("WB-012", 
         "Pengujian logout menghapus seluruh data autentikasi dari localStorage dan redirect ke halaman login.", 
         "localStorage.removeItem('token');\nlocalStorage.removeItem('username');\nlocalStorage.removeItem('role');\nlocalStorage.removeItem('refresh_token');\nwindow.location.href = '/login';", 
         "localStorage.setItem('token', 'some-token');\nlocalStorage.setItem('role', 'pelajar');\nlocalStorage.setItem('username', 'budi');\nAuthService.logout();\nexpect(localStorage.getItem('token')).toBeNull();\nexpect(localStorage.getItem('role')).toBeNull();\nexpect(localStorage.getItem('username')).toBeNull();\nexpect(window.location.href).toBe('/login');", 
         "PASS"),

        # AuthService - getAuthHeaders() & refreshToken() (2 cases)
        ("WB-013", 
         "Pengujian getAuthHeaders mengembalikan Authorization header Bearer token aktif.", 
         "let token = AuthService.getToken();\nif (AuthService.isTokenExpired(token)) {\n  token = await AuthService.refreshToken();\n}\nreturn { Authorization: `Bearer ${token}` };", 
         "const activePayload = { exp: Math.floor(Date.now() / 1000) + 3600 };\nconst token = `header.${btoa(JSON.stringify(activePayload))}.signature`;\nlocalStorage.setItem('token', token);\nconst headers = await AuthService.getAuthHeaders();\nexpect(headers).toEqual({ Authorization: `Bearer ${token}` });", 
         "PASS"),

        ("WB-014", 
         "Pengujian getAuthHeaders memicu refresh token saat token yang tersimpan expired.", 
         "if (AuthService.isTokenExpired(token)) {\n  token = await AuthService.refreshToken();\n}", 
         "const expiredPayload = { exp: Math.floor(Date.now() / 1000) - 3600 };\nconst expiredToken = `header.${btoa(JSON.stringify(expiredPayload))}.signature`;\nlocalStorage.setItem('token', expiredToken);\nlocalStorage.setItem('refresh_token', 'refresh-token');\naxios.post.mockResolvedValueOnce({ data: { token: 'new-jwt-token' } });\nconst headers = await AuthService.getAuthHeaders();\nexpect(localStorage.getItem('token')).toBe('new-jwt-token');\nexpect(headers).toEqual({ Authorization: `Bearer new-jwt-token` });", 
         "PASS"),

        # assessment.js (3 cases)
        ("WB-015", 
         "Pengujian assessment merender tampilan loading sebelum pertanyaan selesai diambil.", 
         "if (loading) return <p>Loading questions...</p>;", 
         "axios.get.mockImplementationOnce(() => new Promise(() => {}));\nrender(<Ass />);\nexpect(screen.getByText('Loading questions...')).toBeInTheDocument();", 
         "PASS"),

        ("WB-016", 
         "Pengujian assessment berhasil mem-fetch dan merender daftar pertanyaan emosi/kondisi dari API.", 
         "const res = await axios.get(`${process.env.NEXT_PUBLIC_API_URL}/api/assessment/questions`, { ... });\nsetQuestions(res.data);", 
         "axios.get.mockResolvedValueOnce({ data: mockQuestions });\nrender(<Ass />);\nawait waitFor(() => { expect(screen.queryByText('Loading questions...')).not.toBeInTheDocument(); });\nexpect(screen.getByText('Stres berlebihan')).toBeInTheDocument();", 
         "PASS"),

        ("WB-017", 
         "Pengujian submit assessment menyertakan intensitas terpilih hanya untuk emosi yang checked=true dan mengarahkan ke halaman rekomendasi.", 
         "const submitted = {\n    answers: Object.entries(answers)\n        .filter(([id, val]) => val.checked)\n        .map(([id, val]) => {\n            const question = questions.find(q => q.id === Number(id));\n            return { code: question?.code, intensity: val.intensity };\n        })\n};\nawait axios.post(`${process.env.NEXT_PUBLIC_API_URL}/api/assessment/submit`, submitted);\nrouter.push(`/home/recommendation?answers=${query}`);", 
         "fireEvent.click(checkboxes[0]);\nfireEvent.change(selects[0], { target: { value: 'high' } });\nfireEvent.click(submitBtn);\nawait waitFor(() => {\n  expect(axios.post).toHaveBeenCalledWith(expect.anything(), { answers: [{ code: 'STRESS', intensity: 'high' }] });\n});\nexpect(mockPush).toHaveBeenCalledWith(expect.stringContaining('/home/recommendation?answers='));", 
         "PASS"),

        # register.js - image upload (4 cases)
        ("WB-018", 
         "Pengujian handleFilePictureChange mendeteksi file foto profil kosong saat registrasi.", 
         "const file = e.target.files ? e.target.files[0] : null;\nif (!file) {\n  setState((prevState) => ({ ...prevState, imageError: \"Tambahkan foto profil.\" }));\n  return;\n}", 
         "const fileInput = screen.getByLabelText(/Tetapkan Foto Profil/i);\nfireEvent.change(fileInput, { target: { files: [] } });\nexpect(screen.getByText('Tambahkan foto profil.')).toBeInTheDocument();", 
         "PASS"),

        ("WB-019", 
         "Pengujian imageLoading diaktifkan sesaat sebelum request upload gambar dikirim ke API.", 
         "setState((prevState) => ({ ...prevState, selectedFile: file, imageLoading: true }));", 
         "axios.post.mockImplementationOnce(() => new Promise(() => {}));\nconst fileInput = screen.getByLabelText(/Tetapkan Foto Profil/i);\nconst file = new File(['mock content'], 'avatar.png', { type: 'image/png' });\nfireEvent.change(fileInput, { target: { files: [file] } });\nexpect(screen.getByText('Uploading image...')).toBeInTheDocument();", 
         "PASS"),

        ("WB-020", 
         "Pengujian penanganan upload image sukses dan menyimpan secure URL gambar.", 
         "const resData = response.data;\nsetState((prev) => ({ ...prev, profile_picture : resData.secure_url || resData.url }));", 
         "axios.post.mockResolvedValueOnce({ data: { secure_url: 'https://cdn.sejiwa.com/avatar.png' } });\nconst fileInput = screen.getByLabelText(/Tetapkan Foto Profil/i);\nconst file = new File(['mock content'], 'avatar.png', { type: 'image/png' });\nfireEvent.change(fileInput, { target: { files: [file] } });\nawait waitFor(() => { expect(screen.getByText('Image uploaded successfully!')).toBeInTheDocument(); });", 
         "PASS"),

        ("WB-021", 
         "Pengujian kegagalan upload image memperbarui state imageError dengan respon API.", 
         "catch (err) {\n  setState((prevState) => ({\n    ...prevState,\n    imageError: err?.response?.data?.error || \"Gagal mengupload gambar. Coba lagi.\",\n  }));\n}", 
         "axios.post.mockRejectedValueOnce({ response: { data: { error: 'File size too large' } } });\nconst fileInput = screen.getByLabelText(/Tetapkan Foto Profil/i);\nconst file = new File(['mock content'], 'avatar.png', { type: 'image/png' });\nfireEvent.change(fileInput, { target: { files: [file] } });\nawait waitFor(() => { expect(screen.getByText('File size too large')).toBeInTheDocument(); });", 
         "PASS"),

        # login.js - redirect logic (3 cases)
        ("WB-022", 
         "Pengujian role-based redirect setelah login sukses (Role Admin).", 
         "if (role === \"admin\") {\n  router.push(\"/admin\");\n}", 
         "AuthService.login.mockResolvedValueOnce({ token: 'mock-admin-token' });\njwtDecode.mockReturnValueOnce({ role: 'admin', id: '1', username: 'admin_user' });\nawait performLogin('admin@sejiwa.com', 'admin123');\nawait waitFor(() => { expect(mockPush).toHaveBeenCalledWith('/admin'); });", 
         "PASS"),

        ("WB-023", 
         "Pengujian role-based redirect setelah login sukses (Role Konselor).", 
         "else if (role === \"konselor\") {\n  router.push(\"/konselor/\");\n}", 
         "AuthService.login.mockResolvedValueOnce({ token: 'mock-konselor-token' });\njwtDecode.mockReturnValueOnce({ role: 'konselor', id: '5', username: 'counselor_user' });\nawait performLogin('counselor@sejiwa.com', 'counselor123');\nawait waitFor(() => { expect(mockPush).toHaveBeenCalledWith('/konselor/'); });", 
         "PASS"),

        ("WB-024", 
         "Pengujian role-based redirect setelah login sukses (Role Pelajar/Default).", 
         "else {\n  router.push(\"/home\");\n}", 
         "AuthService.login.mockResolvedValueOnce({ token: 'mock-pelajar-token' });\njwtDecode.mockReturnValueOnce({ role: 'pelajar', id: '10', username: 'student_user' });\nawait performLogin('student@sejiwa.com', 'student123');\nawait waitFor(() => { expect(mockPush).toHaveBeenCalledWith('/home'); });", 
         "PASS"),

        # home/layout.js - checkToken guard (5 cases)
        ("WB-025", 
         "Pengujian checkToken mendeteksi token tidak ada di localStorage dan mengarahkan ke halaman unauthorized.", 
         "const token = localStorage.getItem('token');\nif (!token) {\n  router.replace('/unauthorized');\n  return;\n}", 
         "render(<Homepage><div data-testid=\"child\">Dashboard Content</div></Homepage>);\nexpect(mockReplace).toHaveBeenCalledWith('/unauthorized');\nexpect(screen.getByText('Loading...')).toBeInTheDocument();", 
         "PASS"),

        ("WB-026", 
         "Pengujian checkToken membiarkan akses dan merender halaman dashboard jika token valid.", 
         "const payload  = JSON.parse(atob(token.split('.')[1]));\nconst expired  = payload.exp * 1000 < Date.now();\nif (expired) { ... } else { setAuthorized(true); }", 
         "const futurePayload = { exp: Math.floor(Date.now() / 1000) + 3600 };\nconst validToken = `header.${btoa(JSON.stringify(futurePayload))}.signature`;\nlocalStorage.setItem('token', validToken);\nrender(<Homepage><div data-testid=\"child\">Dashboard Content</div></Homepage>);\nexpect(screen.getByTestId('child')).toBeInTheDocument();", 
         "PASS"),

        ("WB-027", 
         "Pengujian checkToken memicu logout dan toast error saat mendeteksi token telah kedaluwarsa.", 
         "if (expired) {\n  toast.error('Sesi Anda telah berakhir. Silakan login ulang.');\n  AuthService.logout();\n}", 
         "const expiredPayload = { exp: Math.floor(Date.now() / 1000) - 3600 };\nconst expiredToken = `header.${btoa(JSON.stringify(expiredPayload))}.signature`;\nlocalStorage.setItem('token', expiredToken);\nrender(<Homepage><div data-testid=\"child\">Dashboard Content</div></Homepage>);\nexpect(toast.error).toHaveBeenCalledWith('Sesi Anda telah berakhir. Silakan login ulang.');\nexpect(AuthService.logout).toHaveBeenCalled();", 
         "PASS"),

        ("WB-028", 
         "Pengujian checkToken memicu logout jika format token yang tersimpan di localStorage tidak valid/rusak.", 
         "} catch (err) {\n  console.error('Invalid token :', err);\n  AuthService.logout();\n}", 
         "localStorage.setItem('token', 'invalid-token-format');\nrender(<Homepage><div data-testid=\"child\">Dashboard Content</div></Homepage>);\nexpect(AuthService.logout).toHaveBeenCalled();", 
         "PASS"),

        ("WB-029", 
         "Pengujian checkToken dijalankan secara periodik menggunakan interval waktu 5 detik.", 
         "const interval = setInterval(checkToken, 5000);\nreturn () => clearInterval(interval);", 
         "act(() => { jest.advanceTimersByTime(5000); });\nexpect(AuthService.logout).toHaveBeenCalled();", 
         "PASS"),

        # kelola-jadwal/page.js - fetchUserSchedules filter (2 cases)
        ("WB-030", 
         "Pengujian tombol buat jadwal memicu perpindahan rute ke form pembuatan jadwal konselor.", 
         "const redirectToCreateSchedule = () => router.push('/konselor/kelola-jadwal/create-jadwal');\n// button onClick={redirectToCreateSchedule}", 
         "render(<KelolaJadwal />);\nconst createBtn = screen.getByRole('button', { name: /Buat Jadwal/i });\ncreateBtn.click();\nexpect(mockPush).toHaveBeenCalledWith('/konselor/kelola-jadwal/create-jadwal');", 
         "PASS"),

        ("WB-031", 
         "Pengujian filter data jadwal agar hanya menampilkan jadwal milik counselorId yang sedang login.", 
         "const filtered = allSchedules.filter(\n    (s) => String(s.counselor_id) === String(counselorId)\n);", 
         "axios.get.mockResolvedValueOnce({ data: mockSchedules });\nrender(<KelolaJadwal />);\nawait waitFor(() => {\n  expect(screen.getByText('10:00 - 11:00')).toBeInTheDocument();\n  expect(screen.queryByText('14:00 - 15:00')).not.toBeInTheDocument();\n});", 
         "PASS"),
    ]
    cases.extend(wb_cases)
    
    # Create Table
    headers = ["Test ID", "Deskripsi Pengujian", "Potongan Kode yang Diuji", "Potongan Kode Uji", "Hasil Pengujian"]
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    
    # Set widths (approximate for 6.5 inches printable area)
    col_widths = [Inches(0.8), Inches(1.5), Inches(2.1), Inches(1.5), Inches(0.8)]
    
    # Header styling
    hdr_cells = table.rows[0].cells
    for i, title_text in enumerate(headers):
        hdr_cells[i].text = title_text
        hdr_cells[i].width = col_widths[i]
        set_cell_background(hdr_cells[i], '0E748C') # sky-700
        set_cell_margins(hdr_cells[i], top=120, bottom=120, left=150, right=150)
        
        # Center align headers & make text bold/white
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.name = 'Arial'
        run.font.size = Pt(9.5)
        
    # Populate rows
    for row_idx, case in enumerate(cases):
        row_cells = table.add_row().cells
        
        # Zebra striping
        bg_color = 'F0FDFA' if row_idx % 2 == 0 else 'FFFFFF' # Alternate light cyan-teal & white
        
        for col_idx, cell_value in enumerate(case):
            row_cells[col_idx].text = str(cell_value)
            row_cells[col_idx].width = col_widths[col_idx]
            set_cell_background(row_cells[col_idx], bg_color)
            set_cell_margins(row_cells[col_idx], top=80, bottom=80, left=100, right=100)
            
            p = row_cells[col_idx].paragraphs[0]
            # Alignment rules
            if col_idx in [0, 4]: # Test ID, Hasil Pengujian
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
            if len(p.runs) > 0:
                run = p.runs[0]
                run.font.name = 'Arial'
                run.font.size = Pt(8.5)
                
                # Make code snippets look like code
                if col_idx in [2, 3]:
                    run.font.name = 'Consolas'
                    run.font.size = Pt(8.0)
                    run.font.color.rgb = RGBColor(55, 65, 81) # charcoal
                
                # Color code status if PASS
                if col_idx == 4 and cell_value == "PASS":
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(16, 185, 129) # Green
                elif col_idx == 0:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(75, 85, 99)
            
            row_cells[col_idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Save document
    doc.save("d:\\kuliah\\Projects\\sejiwa\\sejiwa-frontend\\white_box_test.docx")
    print("White Box document generated successfully!")

if __name__ == "__main__":
    create_white_box_document()

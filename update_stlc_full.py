import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import sys

# ======================= DATA CASES =======================
# 100 Black Box Cases
bb_cases_raw = [
    # 1. LANDING PAGE
    ("BB-001", "Verifikasi pemuatan navbar global", "Buka website Sejiwa, periksa keberadaan logo, menu navigasi, tombol Log in dan Sign up.", "URL Utama (/) dan viewport Desktop", "Semua elemen navbar tampil lengkap, memposisikan fixed di bagian atas dengan warna biru sky-600."),
    ("BB-002", "Verifikasi responsive navbar (Mobile View)", "Ubah ukuran layar menjadi mobile (lebar < 768px), periksa kemunculan tombol hamburger.", "Lebar layar 375px", "Navbar menu disembunyikan, tombol hamburger muncul di kanan atas."),
    ("BB-003", "Verifikasi pembukaan menu laci di Mobile View", "Klik tombol hamburger di resolusi mobile.", "Klik event pada button hamburger", "Drawer/Dialog menu mobile terbuka secara halus menampilkan seluruh item navigasi."),
    ("BB-004", "Verifikasi penutupan menu laci di Mobile View", "Klik tombol 'XMark' atau area luar laci saat menu terbuka.", "Klik event pada button 'XMark'", "Drawer menu menutup kembali dengan benar."),
    ("BB-005", "Verifikasi navigasi link 'Home' pada navbar", "Klik link 'Home' di navbar.", "Klik link 'Home'", "Browser berpindah rute ke halaman '/home'."),
    ("BB-006", "Verifikasi navigasi link 'Assessment' pada navbar", "Klik link 'Assessment' di navbar.", "Klik link 'Assessment'", "Browser berpindah rute ke halaman '/home/assessment'."),
    ("BB-007", "Verifikasi klik tombol 'Log in' di navbar", "Klik tombol 'Log in' di sudut kanan navbar.", "Klik link 'Log in'", "Browser berpindah rute ke halaman '/login'."),
    ("BB-008", "Verifikasi klik tombol 'Sign up' di navbar", "Klik tombol 'Sign up' di sudut kanan navbar.", "Klik link 'Sign up'", "Browser berpindah rute ke halaman '/register'."),
    ("BB-009", "Verifikasi tombol CTA Utama Hero Section", "Klik tombol 'Assessment' di area Hero Section.", "Klik tombol 'Assessment' (Hero)", "Browser mengarahkan ke halaman '/home/assessment'."),
    ("BB-010", "Verifikasi tautan eksternal media sosial di footer", "Scroll ke footer, klik link Facebook dan Instagram.", "Klik link media sosial", "Tab baru terbuka atau mengarah ke URL eksternal media sosial terkait."),
    
    # 2. AUTH - LOGIN
    ("BB-011", "Verifikasi UI halaman login", "Buka halaman '/login', periksa formulir email, password, tombol login, dan link registrasi.", "Navigasi ke '/login'", "UI form login tampil dengan layout grid, gradien sky-700, font Rubik, dan Josefin Slab."),
    ("BB-012", "Login sukses sebagai Pelajar (Role Student)", "Masukkan email dan password pelajar terdaftar, klik Login.", "Email: pelajar@sejiwa.com, Pass: password123", "Menyimpan JWT token & role ke localStorage, lalu redirect ke '/home'."),
    ("BB-013", "Login sukses sebagai Konselor (Role Counselor)", "Masukkan email dan password konselor terdaftar, klik Login.", "Email: konselor@sejiwa.com, Pass: password123", "Menyimpan JWT token & role ke localStorage, redirect ke '/konselor/'."),
    ("BB-014", "Login sukses sebagai Admin (Role Admin)", "Masukkan email dan password admin terdaftar, klik Login.", "Email: admin@sejiwa.com, Pass: password123", "Menyimpan JWT token & role ke localStorage, redirect ke '/admin'."),
    ("BB-015", "Login gagal dengan email tidak terdaftar", "Masukkan email acak dan password salah, klik Login.", "Email: fakeuser@sejiwa.com, Pass: salah123", "Tampil pesan kesalahan error: 'Email atau password salah.'"),
    ("BB-016", "Login gagal dengan password salah", "Masukkan email terdaftar tapi password salah, klik Login.", "Email: pelajar@sejiwa.com, Pass: salah123", "Tampil pesan kesalahan error: 'Email atau password salah.'"),
    ("BB-017", "Validasi format email pada form login", "Masukkan format email tanpa '@' atau domain, klik Login.", "Email: pelajar-invalid, Pass: password123", "Pemberitahuan validasi HTML5 input type='email' wajib diisi format valid."),
    ("BB-018", "Validasi form kosong (Empty Email & Password)", "Biarkan input email dan password kosong, klik Login.", "Email: [Empty], Pass: [Empty]", "HTML5 memunculkan peringatan field required pada input email/password."),
    ("BB-019", "Redirect link register dari login", "Klik teks link 'Register' di bawah form login.", "Klik link Register", "Halaman berpindah langsung ke '/register'."),
    ("BB-020", "Verifikasi penyimpanan JWT di LocalStorage", "Setelah login sukses, periksa Application tab di DevTools.", "Inspect LocalStorage", "Terdapat key 'token', 'refresh_token', 'username', dan 'role' dengan nilai valid."),
    ("BB-021", "Logout dari dashboard Pelajar", "Klik tombol logout di sidebar dashboard pelajar.", "Klik Logout", "Token dihapus dari LocalStorage, user di-redirect ke halaman '/login'."),
    ("BB-022", "Logout dari dashboard Konselor", "Klik tombol logout di sidebar dashboard konselor.", "Klik Logout", "Token dihapus dari LocalStorage, user di-redirect ke halaman '/login'."),
    ("BB-023", "Penyimpanan otomatis username dari email jika payload JWT tidak ada username", "Lakukan login dengan akun tanpa field username di JWT payload.", "Email: andi.kurniawan@sejiwa.com", "Sistem memanggil utilitas generateUsernameFromEmail dan menyimpan 'Andi' di localStorage."),
    ("BB-024", "Pencegahan akses login kembali saat token aktif", "Coba akses halaman '/login' ketika status session terautentikasi.", "Akses URL '/login'", "User diarahkan kembali ke '/home' secara otomatis."),
    ("BB-025", "Verifikasi token expired memicu logout paksa", "Ubah expired date token JWT di LocalStorage menjadi masa lampau.", "Simulasi token expired", "Interval checkToken mendeteksi token expired, memunculkan Toast error, lalu melakukan logout."),
    
    # 3. AUTH - REGISTER
    ("BB-026", "Verifikasi pembagian step registrasi (Step 1)", "Buka '/register', verifikasi form pendaftaran awal.", "Navigasi '/register'", "Tampil formulir Step 1: email, password, confirmPassword, dan pilihan role."),
    ("BB-027", "Verifikasi validasi Step 1 kosong", "Klik tombol Register pada Step 1 tanpa mengisi data apapun.", "Klik Register (Step 1)", "Field formulir ditolak dan muncul indikasi wajib isi."),
    ("BB-028", "Validasi perbedaan password dan konfirmasi password", "Isi password 'abc12345' dan konfirmasi password 'xyz98765' pada Step 1.", "Pass: abc12345, Confirm: xyz98765", "Tampil pesan kesalahan: 'Ketik ulang password dengan benar.' setelah step 2 di-submit."),
    ("BB-029", "Pemberitahuan role belum terpilih di Step 1", "Isi semua data Step 1 kecuali dropdown Role, klik Register.", "Role: Pilih Role (default)", "Registrasi tidak berlanjut atau menampilkan validasi role wajib dipilih."),
    ("BB-030", "Navigasi ke Step 2 setelah Step 1 valid", "Isi lengkap dan benar Step 1, lalu klik tombol Register.", "Email, Pass, Confirm Pass valid, Role Pelajar", "Formulir beralih ke Step 2 (pengunggahan foto profil dan input username)."),
    ("BB-031", "Tombol 'Back' di Step 2 mengembalikan data Step 1", "Di Step 2, klik tombol 'Back'.", "Klik Back", "Formulir kembali ke Step 1 dengan data email & password yang sudah diinput sebelumnya tetap ada."),
    ("BB-032", "Verifikasi upload foto profil dengan file gambar valid", "Di Step 2, pilih file gambar PNG/JPG ukuran < 2MB.", "Pilih file 'profile.jpg'", "Menampilkan teks 'Uploading image...', respons sukses secure_url, dan preview foto bulat."),
    ("BB-033", "Upload foto profil gagal format file tidak valid", "Pilih file non-gambar (misal file PDF) pada input profile picture.", "Pilih file 'document.pdf'", "Muncul pesan kesalahan error: 'Gagal mengupload gambar. Coba lagi.' dari API upload."),
    ("BB-034", "Validasi submit Step 2 tanpa upload foto profil", "Isi username namun kosongkan foto profil, klik Register.", "Username: budi123, Profile Pic: [None]", "Muncul pesan error: 'Harap mengisi semua form. Pastikan Anda sudah mengisi semua form dan memasang foto profil'"),
    ("BB-035", "Validasi submit Step 2 tanpa username", "Upload foto profil sukses namun kosongkan username, klik Register.", "Username: [Empty], Profile Pic: uploaded", "Input username ditandai required oleh HTML5 validation."),
    ("BB-036", "Submit registrasi sukses sebagai Pelajar", "Selesaikan Step 1 & 2 dengan benar untuk role Pelajar, klik Register.", "Role: Pelajar, Username: student1, Image valid", "Data berhasil dikirim ke API, menampilkan pesan sukses, menyimpan token, dan redirect ke '/login'."),
    ("BB-037", "Submit registrasi sukses sebagai Konselor", "Selesaikan Step 1 & 2 dengan benar untuk role Konselor, klik Register.", "Role: Konselor, Username: counselor1, Image valid", "Data berhasil dikirim ke API, menampilkan pesan sukses, menyimpan token, dan redirect ke '/login'."),
    ("BB-038", "Penanganan error server saat registrasi (misal email duplikat)", "Gunakan email yang sudah terdaftar untuk melakukan registrasi baru.", "Email: terdaftar@sejiwa.com", "Menampilkan pesan error dari backend: 'Email sudah terdaftar' atau sejenisnya."),
    ("BB-039", "Verifikasi loading state saat submit registrasi", "Klik register saat proses upload foto sedang berjalan.", "Proses upload gambar aktif", "Tombol submit dinonaktifkan (disabled) dan berlabel 'Uploading Image...'."),
    ("BB-040", "Redirect ke login via link footer register", "Klik link 'Login' di bagian bawah formulir registrasi.", "Klik link Login", "Browser berpindah rute ke halaman '/login'."),
    
    # 4. HOME & LAYOUT GUARDS
    ("BB-041", "Akses halaman '/home' tanpa token (Guard check)", "Hapus localStorage token, lalu coba buka rute '/home'.", "Akses URL '/home' tanpa token", "Sistem langsung mengarahkan (replace) rute ke '/unauthorized'."),
    ("BB-042", "Akses halaman '/home' dengan token valid", "Login sukses, pastikan token tersimpan, muat halaman '/home'.", "Token valid terdeteksi", "Halaman dimuat sempurna menampilkan layout dashboard (Navbar, Sidebar, Content)."),
    ("BB-043", "Verifikasi Sidebar toggle menu", "Klik tombol menu (toggle) di navbar.", "Klik Toggle Sidebar", "Sidebar bergeser masuk (buka) / keluar (tutup) dari layar dengan lancar."),
    ("BB-044", "Verifikasi link Navigasi Sidebar - Home", "Klik menu 'Home' di sidebar.", "Klik 'Home'", "Rute diarahkan ke '/home' tanpa reload penuh."),
    ("BB-045", "Verifikasi link Navigasi Sidebar - Assessment", "Klik menu 'Assessment' di sidebar.", "Klik 'Assessment'", "Rute diarahkan ke '/home/assessment'."),
    ("BB-046", "Verifikasi link Navigasi Sidebar - Chat", "Klik menu 'Chat' di sidebar.", "Klik 'Chat'", "Rute diarahkan ke '/home/chat'."),
    ("BB-047", "Verifikasi link Navigasi Sidebar - Rekomendasi", "Klik menu 'Rekomendasi' di sidebar.", "Klik 'Rekomendasi'", "Rute diarahkan ke '/home/recommendation'."),
    ("BB-048", "Verifikasi link Navigasi Sidebar - Settings", "Klik menu 'Settings' di sidebar.", "Klik 'Settings'", "Rute diarahkan ke '/home/settings'."),
    ("BB-049", "Pengecekan token berkala setiap 5 detik", "Buka '/home', tunggu lebih dari 5 detik, periksa console log / network.", "Biarkan halaman aktif > 5 detik", "Fungsi checkToken terpanggil secara berulang (setinterval) mendeteksi status keaktifan sesi."),
    ("BB-050", "Verifikasi detail profil pengguna di Navbar", "Periksa tampilan username yang tersimpan di localStorage pada navbar.", "Login dengan user 'Budi'", "Navbar menampilkan teks sapaan / username 'Budi' dengan benar."),

    # 5. ASSESSMENT
    ("BB-051", "Verifikasi pengambilan daftar pertanyaan assessment", "Buka '/home/assessment', periksa apakah daftar emosi/kondisi berhasil di-fetch.", "Navigasi ke '/home/assessment'", "Daftar emosi terisi lengkap sesuai response API `/api/assessment/questions`."),
    ("BB-052", "Tampilan loading saat mengambil pertanyaan", "Simulasikan koneksi lambat, buka halaman assessment.", "Koneksi lambat", "Muncul teks 'Loading questions...' sebelum tabel assessment dirender."),
    ("BB-053", "Verifikasi kolom tabel assessment", "Buka assessment, periksa header kolom.", "Lihat tabel", "Tabel memiliki kolom: Emosi/Kondisi, Check, Intensitas, dan Kata-kata Motivasi."),
    ("BB-054", "Mencentang checkbox emosi/kondisi", "Klik checkbox pada salah satu baris emosi.", "Check emosi 'Stres'", "Checkbox tercentang, mengubah state checked untuk emosi tersebut menjadi true."),
    ("BB-055", "Memilih intensitas emosi (Low/Medium/High)", "Pilih opsi dari dropdown intensitas di baris emosi.", "Pilih 'Medium' pada emosi 'Cemas'", "Nilai state intensitas terupdate dengan 'medium'."),
    ("BB-056", "Menghapus centang pada emosi yang dipilih", "Klik kembali checkbox yang sudah tercentang.", "Uncheck emosi", "Checkbox kosong, mengubah state checked menjadi false."),
    ("BB-057", "Fitur pencarian daftar emosi (Search Input)", "Ketik nama emosi tertentu pada input pencarian.", "Ketik 'Sedih'", "Tabel menyaring baris sehingga hanya menampilkan emosi yang mengandung kata 'Sedih'."),
    ("BB-058", "Verifikasi pagination halaman pertanyaan (Next Button)", "Klik tombol 'Next' di bagian bawah tabel.", "Klik Next", "Tabel memperbarui data menampilkan pertanyaan halaman berikutnya."),
    ("BB-059", "Verifikasi pagination halaman pertanyaan (Prev Button)", "Klik tombol 'Prev' setelah berada di page 2.", "Klik Prev", "Tabel kembali menampilkan halaman sebelumnya."),
    ("BB-060", "Pengubahan opsi jumlah baris data per halaman (Rows per page)", "Ubah dropdown jumlah data di footer (misal ke '3').", "Pilih opsi '3'", "Tabel membatasi baris data yang ditampilkan maksimal 3 per halaman."),
    ("BB-061", "Submit assessment dengan data valid", "Pilih minimal 1 emosi dengan intensitasnya, klik 'Kirim Assessment'.", "Submit payload valid", "Mengirim payload POST ke `/api/assessment/submit` dan redirect ke rekomendasi dengan query string."),
    ("BB-062", "Submit assessment tanpa memilih emosi apapun", "Klik tombol 'Kirim Assessment' langsung tanpa mencentang emosi.", "Klik Kirim Assessment", "Payload terkirim dengan array answers kosong `[]` atau dicegah oleh sistem front-end."),
    ("BB-063", "Navigasi ke halaman Riwayat Assessment", "Klik tombol 'Hasil Assessment' di kanan atas halaman.", "Klik Hasil Assessment", "Browser berpindah ke rute '/home/assessment/hasil-assessment'."),
    ("BB-064", "Pemuatan halaman Hasil Assessment", "Buka rute '/home/assessment/hasil-assessment', verifikasi data.", "Navigasi hasil-assessment", "Menampilkan riwayat assessment pengguna dalam bentuk tabel atau list ringkasan."),
    ("BB-065", "Penanganan error saat submit assessment gagal", "Simulasi kegagalan koneksi API post submit, klik kirim.", "Koneksi terputus", "Muncul pop-up alert bertuliskan: 'Failed to submit.'"),

    # 6. RECOMMENDATION
    ("BB-066", "Pemuatan halaman rekomendasi dari query parameter", "Buka '/home/recommendation?answers=[...]', verifikasi konten.", "Akses URL dengan param answers", "Halaman berhasil memparsing jawaban dan memuat data artikel/video rekomendasi."),
    ("BB-067", "Tampilan daftar video rekomendasi", "Buka rekomendasi, verifikasi komponen video.", "Lihat bagian video", "Menampilkan list thumbnail video kesehatan mental beserta judul."),
    ("BB-068", "Tampilan daftar artikel rekomendasi", "Buka rekomendasi, verifikasi komponen artikel.", "Lihat bagian artikel", "Menampilkan list ringkasan artikel terkait emosi yang dipilih."),
    ("BB-069", "Navigasi klik video rekomendasi", "Klik salah satu kartu video rekomendasi.", "Klik Video", "Membuka video di tab baru atau memutar video langsung."),
    ("BB-070", "Navigasi klik artikel rekomendasi", "Klik salah satu tautan artikel rekomendasi.", "Klik Artikel", "Membuka halaman detail artikel atau rute eksternal sumber artikel."),

    # 7. CHAT & FIND COUNSELOR
    ("BB-071", "Akses dashboard chat '/home/chat'", "Buka halaman utama chat, periksa menu grid.", "Navigasi '/home/chat'", "Menampilkan grid opsi: Cari Konselor, Chat Room, dan Booking."),
    ("BB-072", "Navigasi ke menu 'Cari Konselor'", "Klik kartu 'Cari Konselor' di dashboard chat.", "Klik Cari Konselor", "Browser berpindah ke rute '/home/chat/find-conselor'."),
    ("BB-073", "Tampilan halaman 'Cari Konselor' (find-conselor)", "Buka '/home/chat/find-conselor', periksa daftar konselor.", "Navigasi find-conselor", "Menampilkan daftar konselor yang berstatus aktif beserta spesialisasi."),
    ("BB-074", "Navigasi ke menu 'Chat Room'", "Klik kartu 'Chat Room' di dashboard chat.", "Klik Chat Room", "Browser berpindah ke rute '/home/chat/chat-pelajar'."),
    ("BB-075", "Tampilan halaman 'Chat Room' (chat-pelajar)", "Buka '/home/chat/chat-pelajar', periksa antarmuka chat.", "Navigasi chat-pelajar", "Menampilkan riwayat percakapan di panel kiri dan jendela pesan aktif di kanan."),
    ("BB-076", "Mengirim pesan teks di Chat Room", "Ketik pesan di kolom input chat, klik tombol kirim.", "Ketik 'Halo konselor', klik send", "Pesan baru terkirim, tampil di bubble chat kanan, terkirim via websocket."),
    ("BB-077", "Menerima pesan real-time dari konselor", "Konselor mengirimkan pesan balasan dari dashboard konselor.", "Balasan dari konselor", "Pesan langsung muncul di bubble chat kiri secara real-time tanpa refresh."),
    ("BB-078", "Mengirim pesan kosong di Chat Room", "Biarkan input chat kosong, klik tombol kirim.", "Input kosong, klik send", "Aksi kirim diabaikan, tidak ada bubble chat kosong yang ditambahkan."),
    ("BB-079", "Status koneksi socket di Chat Room", "Buka halaman chat, cek konektivitas websocket.", "Koneksi socket.io", "Menghubungkan ke socket server `/api/socket` dengan status connected."),
    ("BB-080", "Scroll otomatis ke pesan terbawah", "Kirim/terima banyak pesan hingga melebihi tinggi layar chat.", "Kirim > 10 pesan", "Jendela percakapan otomatis scroll ke bagian pesan paling bawah/terbaru."),

    # 8. BOOKING SYSTEM
    ("BB-081", "Akses menu 'Booking' pelajar", "Klik kartu 'Booking' di dashboard chat.", "Klik Booking", "Browser berpindah ke rute '/home/chat/booking'."),
    ("BB-082", "Menampilkan pilihan jadwal konselor untuk dibooking", "Buka form pemesanan sesi konseling (create-booking).", "Akses create-booking", "Menampilkan list tanggal & jam ketersediaan jadwal konselor yang aktif."),
    ("BB-083", "Membuat booking baru dengan memilih jadwal valid", "Pilih salah satu jadwal konselor, klik tombol Konfirmasi Booking.", "Pilih jadwal, klik booking", "Mengirim data ke backend, mengarahkan ke halaman success-booking."),
    ("BB-084", "Halaman sukses booking (success-booking)", "Verifikasi pemuatan halaman setelah booking berhasil dibuat.", "Redirect success-booking", "Tampil teks sukses pemesanan sesi konseling dan tombol kembali."),
    ("BB-085", "Melihat riwayat booking pelajar (history-booking)", "Buka halaman history-booking.", "Navigasi history-booking", "Menampilkan tabel riwayat booking: nama konselor, tanggal, jam, dan status (pending/approved/rejected)."),
    ("BB-086", "Booking dengan jadwal yang tidak tersedia", "Coba booking jadwal yang is_available-nya false.", "Pilih jadwal tidak tersedia", "Aksi ditolak, tombol booking dinonaktifkan atau memunculkan error validation."),
    ("BB-087", "Akses menu kelola booking oleh konselor", "Login sebagai konselor, masuk ke rute '/konselor/bookings'.", "Navigasi /konselor/bookings", "Menampilkan daftar request booking dari pelajar yang berstatus Pending."),
    ("BB-088", "Konselor menyetujui request booking (Approve)", "Klik tombol 'Approve' pada salah satu request booking pending.", "Klik Approve", "Status booking berubah menjadi 'approved' / 'confirm-booking', pelajar mendapat notifikasi."),
    ("BB-089", "Konselor menolak request booking (Reject)", "Klik tombol 'Reject' pada salah satu request booking pending.", "Klik Reject", "Status booking berubah menjadi 'rejected-booking', jadwal kembali tersedia."),
    ("BB-090", "Perubahan status booking real-time pada riwayat pelajar", "Pelajar memeriksa history-booking setelah konselor menyetujui sesi.", "Periksa status booking", "Status booking di baris terkait terupdate menjadi 'Approved' / 'Terkonfirmasi'."),

    # 9. COUNSELOR - KELOLA JADWAL
    ("BB-091", "Akses menu 'Kelola Jadwal' konselor", "Login sebagai konselor, klik 'Kelola Jadwal' di sidebar.", "Navigasi /konselor/kelola-jadwal", "Halaman berhasil memuat jadwal pribadi milik konselor bersangkutan (filter by counselor_id)."),
    ("BB-092", "Menambahkan jadwal baru (create-jadwal)", "Klik 'Buat Jadwal', isi tanggal, jam, dan status ketersediaan, klik Simpan.", "Tanggal: 15-07-2026, Jam: 10:00, Tersedia: True", "Jadwal baru tersimpan di database, muncul pesan 'Schedule successfully saved!' dan list ter-refresh."),
    ("BB-093", "Mengubah jadwal ketersediaan (update-jadwal)", "Klik 'Ubah Jadwal', ubah status ketersediaan menjadi Tidak Tersedia.", "Ubah is_available ke False", "Data jadwal berhasil diperbarui di database dan tampilan berubah menjadi Tidak Tersedia."),
    ("BB-094", "Validasi form tambah jadwal kosong", "Klik Simpan pada form tambah jadwal tanpa memilih tanggal / jam.", "Klik Simpan (kosong)", "Tampil pesan kesalahan 'Failed to save schedule.' atau peringatan input kosong."),
    ("BB-095", "Filter otomatis jadwal hanya milik konselor yang login", "Periksa daftar jadwal yang tampil di dashboard.", "Login Counselor ID: 5", "Semua jadwal yang terdaftar memiliki counselor_id bernilai 5 saja, jadwal konselor lain tidak tampil."),

    # 10. ADVANCED ROLES & PERMISSIONS GUARD
    ("BB-096", "Pelajar mencoba mengakses rute admin (/admin)", "Login sebagai pelajar, lalu ketik manual URL '/admin' di browser.", "Akses URL '/admin' dengan role pelajar", "Sistem memblokir akses dan mengarahkan ke halaman '/unauthorized'."),
    ("BB-097", "Pelajar mencoba mengakses rute konselor (/konselor)", "Login sebagai pelajar, lalu ketik manual URL '/konselor' di browser.", "Akses URL '/konselor' dengan role pelajar", "Sistem memblokir akses dan mengarahkan ke halaman '/unauthorized'."),
    ("BB-098", "Konselor mencoba mengakses rute admin (/admin)", "Login sebagai konselor, ketik manual URL '/admin' di browser.", "Akses URL '/admin' dengan role konselor", "Sistem memblokir akses dan mengarahkan ke halaman '/unauthorized'."),
    ("BB-099", "Guest/User tidak terautentikasi mencoba mengakses halaman dashboard", "Hapus token, buka '/home/settings' secara langsung.", "Akses URL '/home/settings' tanpa token", "Sistem mendeteksi tiadanya token dan mengarahkan ke '/unauthorized'."),
    ("BB-100", "Akses halaman '/unauthorized'", "Buka URL '/unauthorized' secara langsung.", "Navigasi '/unauthorized'", "Menampilkan pesan peringatan hak akses ditolak dan menyediakan tombol kembali ke Login.")
]

# 31 White Box Cases
wb_cases = [
    # generateUsernameFromEmail.js (5 cases)
    ("WB-001", "Pengujian utilitas generateUsernameFromEmail dengan input email normal.", "const localPart = email.split('@')[0];\nconst nameCandidate = localPart.split(/[.\\d_]/)[0];\nreturn nameCandidate.charAt(0).toUpperCase() + nameCandidate.slice(1);", "expect(generateUsernameFromEmail('budi.kurniawan123@gmail.com')).toBe('Budi');"),
    ("WB-002", "Pengujian utilitas generateUsernameFromEmail dengan email tanpa domain.", "if (!email) return \"User\";\nconst localPart = email.split('@')[0];", "expect(generateUsernameFromEmail('riri')).toBe('Riri');"),
    ("WB-003", "Pengujian utilitas generateUsernameFromEmail dengan email kosong/null/undefined.", "if (!email) return \"User\";", "expect(generateUsernameFromEmail('')).toBe('User');\nexpect(generateUsernameFromEmail(null)).toBe('User');\nexpect(generateUsernameFromEmail(undefined)).toBe('User');"),
    ("WB-004", "Pengujian utilitas generateUsernameFromEmail dengan email berawalan angka atau karakter khusus.", "const nameCandidate = localPart.split(/[.\\d_]/)[0];\nif (!nameCandidate) return \"User\";", "expect(generateUsernameFromEmail('123budi@gmail.com')).toBe('User');\nexpect(generateUsernameFromEmail('.andi@gmail.com')).toBe('User');"),
    ("WB-005", "Pengujian utilitas generateUsernameFromEmail dengan email yang mengandung angka atau tanda titik/garis bawah di tengah.", "const nameCandidate = localPart.split(/[.\\d_]/)[0];", "expect(generateUsernameFromEmail('andi_kurniawan@sejiwa.com')).toBe('Andi');\nexpect(generateUsernameFromEmail('rudi123@sejiwa.com')).toBe('Rudi');"),

    # AuthService - isTokenExpired() (4 cases)
    ("WB-006", "Pengujian isTokenExpired dengan format token tidak memiliki 3 bagian (dot split).", "if (!token || token.split('.').length !== 3) { return true; }", "expect(AuthService.isTokenExpired('')).toBe(true);\nexpect(AuthService.isTokenExpired('invalid-token')).toBe(true);"),
    ("WB-007", "Pengujian isTokenExpired dengan token yang sudah kedaluwarsa.", "const payload = JSON.parse(atob(token.split('.')[1]));\nreturn payload.exp * 1000 < Date.now();", "const expiredPayload = { exp: Math.floor(Date.now() / 1000) - 3600 };\nconst token = `header.${btoa(JSON.stringify(expiredPayload))}.signature`;\nexpect(AuthService.isTokenExpired(token)).toBe(true);"),
    ("WB-008", "Pengujian isTokenExpired dengan token aktif yang belum kedaluwarsa.", "const payload = JSON.parse(atob(token.split('.')[1]));\nreturn payload.exp * 1000 < Date.now();", "const futurePayload = { exp: Math.floor(Date.now() / 1000) + 3600 };\nconst token = `header.${btoa(JSON.stringify(futurePayload))}.signature`;\nexpect(AuthService.isTokenExpired(token)).toBe(false);"),
    ("WB-009", "Pengujian isTokenExpired saat token rusak pada bagian base64 payload.", "try {\n  const payload = JSON.parse(atob(token.split('.')[1]));\n} catch (error) {\n  return true;\n}", "const brokenToken = 'header.brokenbase64payload.signature';\nexpect(AuthService.isTokenExpired(brokenToken)).toBe(true);"),

    # AuthService - login() (2 cases)
    ("WB-010", "Pengujian login sukses menyimpan token JWT dan refresh token ke localStorage.", "localStorage.setItem('token', jwtToken);\nlocalStorage.setItem('refresh_token', refreshToken);", "axios.post.mockResolvedValueOnce({ data: { token: 'mock-jwt', refresh_token: 'mock-rt' } });"),
    ("WB-011", "Pengujian login gagal melempar error dengan pesan kegagalan dari server.", "catch (error) {\n  throw new Error(error.response?.data?.message || 'Login failed');\n}", "axios.post.mockRejectedValueOnce({ response: { data: { message: 'Wrong password' } } });"),

    # AuthService - logout() (1 case)
    ("WB-012", "Pengujian logout menghapus seluruh data autentikasi dari localStorage dan redirect ke halaman login.", "localStorage.removeItem('token');\nwindow.location.href = '/login';", "AuthService.logout();\nexpect(localStorage.getItem('token')).toBeNull();"),

    # AuthService - getAuthHeaders() & refreshToken() (2 cases)
    ("WB-013", "Pengujian getAuthHeaders mengembalikan Authorization header Bearer token aktif.", "return { Authorization: `Bearer ${token}` };", "const headers = await AuthService.getAuthHeaders();\nexpect(headers).toEqual({ Authorization: `Bearer ${token}` });"),
    ("WB-014", "Pengujian getAuthHeaders memicu refresh token saat token yang tersimpan expired.", "if (AuthService.isTokenExpired(token)) {\n  token = await AuthService.refreshToken();\n}", "localStorage.setItem('token', expiredToken);\nconst headers = await AuthService.getAuthHeaders();\nexpect(localStorage.getItem('token')).toBe('new-jwt-token');"),

    # assessment.js (3 cases)
    ("WB-015", "Pengujian assessment merender tampilan loading sebelum pertanyaan selesai diambil.", "if (loading) return <p>Loading questions...</p>;", "render(<Ass />);\nexpect(screen.getByText('Loading questions...')).toBeInTheDocument();"),
    ("WB-016", "Pengujian assessment berhasil mem-fetch dan merender daftar pertanyaan emosi dari API.", "const res = await axios.get(`/api/assessment/questions`);\nsetQuestions(res.data);", "render(<Ass />);\nexpect(screen.getByText('Stres berlebihan')).toBeInTheDocument();"),
    ("WB-017", "Pengujian submit assessment menyertakan intensitas terpilih hanya untuk emosi yang checked=true.", "answers: Object.entries(answers).filter(([id, val]) => val.checked)", "await waitFor(() => {\n  expect(axios.post).toHaveBeenCalledWith(expect.anything(), { answers: [{ code: 'STRESS', intensity: 'high' }] });\n});"),

    # register.js - image upload (4 cases)
    ("WB-018", "Pengujian handleFilePictureChange mendeteksi file foto profil kosong saat registrasi.", "if (!file) {\n  setState({ imageError: \"Tambahkan foto profil.\" });\n  return;\n}", "fireEvent.change(fileInput, { target: { files: [] } });\nexpect(screen.getByText('Tambahkan foto profil.')).toBeInTheDocument();"),
    ("WB-019", "Pengujian imageLoading diaktifkan sesaat sebelum request upload gambar dikirim ke API.", "setState({ ...prevState, selectedFile: file, imageLoading: true });", "fireEvent.change(fileInput, { target: { files: [file] } });\nexpect(screen.getByText('Uploading image...')).toBeInTheDocument();"),
    ("WB-020", "Pengujian penanganan upload image sukses dan menyimpan secure URL gambar.", "setState({ profile_picture : resData.secure_url });", "axios.post.mockResolvedValueOnce({ data: { secure_url: 'https://cdn.sejiwa.com/avatar.png' } });"),
    ("WB-021", "Pengujian kegagalan upload image memperbarui state imageError dengan respon API.", "catch (err) {\n  setState({ imageError: err?.response?.data?.error });\n}", "axios.post.mockRejectedValueOnce({ response: { data: { error: 'File size too large' } } });"),

    # login.js - redirect logic (3 cases)
    ("WB-022", "Pengujian role-based redirect setelah login sukses (Role Admin).", "if (role === \"admin\") {\n  router.push(\"/admin\");\n}", "await waitFor(() => { expect(mockPush).toHaveBeenCalledWith('/admin'); });"),
    ("WB-023", "Pengujian role-based redirect setelah login sukses (Role Konselor).", "else if (role === \"konselor\") {\n  router.push(\"/konselor/\");\n}", "await waitFor(() => { expect(mockPush).toHaveBeenCalledWith('/konselor/'); });"),
    ("WB-024", "Pengujian role-based redirect setelah login sukses (Role Pelajar).", "else {\n  router.push(\"/home\");\n}", "await waitFor(() => { expect(mockPush).toHaveBeenCalledWith('/home'); });"),

    # home/layout.js - checkToken guard (5 cases)
    ("WB-025", "Pengujian checkToken mendeteksi token tidak ada di localStorage dan mengarahkan ke halaman unauthorized.", "if (!token) {\n  router.replace('/unauthorized');\n  return;\n}", "render(<Homepage />);\nexpect(mockReplace).toHaveBeenCalledWith('/unauthorized');"),
    ("WB-026", "Pengujian checkToken membiarkan akses dan merender halaman dashboard jika token valid.", "if (expired) { ... } else { setAuthorized(true); }", "render(<Homepage />);\nexpect(screen.getByTestId('child')).toBeInTheDocument();"),
    ("WB-027", "Pengujian checkToken memicu logout dan toast error saat mendeteksi token telah kedaluwarsa.", "if (expired) {\n  toast.error('Sesi Anda telah berakhir.');\n  AuthService.logout();\n}", "render(<Homepage />);\nexpect(AuthService.logout).toHaveBeenCalled();"),
    ("WB-028", "Pengujian checkToken memicu logout jika format token yang tersimpan di localStorage tidak valid/rusak.", "} catch (err) {\n  AuthService.logout();\n}", "localStorage.setItem('token', 'invalid-token-format');\nrender(<Homepage />);\nexpect(AuthService.logout).toHaveBeenCalled();"),
    ("WB-029", "Pengujian checkToken dijalankan secara periodik menggunakan interval waktu 5 detik.", "const interval = setInterval(checkToken, 5000);\nreturn () => clearInterval(interval);", "act(() => { jest.advanceTimersByTime(5000); });\nexpect(AuthService.logout).toHaveBeenCalled();"),

    # kelola-jadwal/page.js - fetchUserSchedules filter (2 cases)
    ("WB-030", "Pengujian tombol buat jadwal memicu perpindahan rute ke form pembuatan jadwal konselor.", "const redirectToCreateSchedule = () => router.push('/konselor/kelola-jadwal/create-jadwal');", "const createBtn = screen.getByRole('button', { name: /Buat Jadwal/i });\ncreateBtn.click();\nexpect(mockPush).toHaveBeenCalledWith('/konselor/kelola-jadwal/create-jadwal');"),
    ("WB-031", "Pengujian filter data jadwal agar hanya menampilkan jadwal milik counselorId yang sedang login.", "const filtered = allSchedules.filter(\n    (s) => String(s.counselor_id) === String(counselorId)\n);", "await waitFor(() => {\n  expect(screen.getByText('10:00 - 11:00')).toBeInTheDocument();\n});")
]

# ======================= UTILS =======================

def set_cell_background(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def insert_paragraph_after(paragraph, text, style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = docx.text.paragraph.Paragraph(new_p, paragraph._parent)
    if text:
        new_para.add_run(text)
    if style:
        new_para.style = style
    return new_para

def delete_paragraph(paragraph):
    p = paragraph._element
    if p.getparent() is not None:
        p.getparent().remove(p)
    paragraph._p = paragraph._element = None

def update_bab_a(doc):
    fr_heading_idx = -1
    nfr_heading_idx = -1
    rtm_heading_idx = -1
    
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip().lower()
        if "functional requirement" in text and ("2." in text or "2 " in text):
            fr_heading_idx = i
        elif "non functional requirement" in text and ("3." in text or "3 " in text):
            nfr_heading_idx = i
        elif "requirements traceability matrix" in text or "rtm" in text.lower().split():
            if "4." in text or "4 " in text:
                rtm_heading_idx = i
                
    if rtm_heading_idx != -1:
        p_rtm = doc.paragraphs[rtm_heading_idx]
        idx = rtm_heading_idx + 1
        while idx < len(doc.paragraphs):
            nxt_p = doc.paragraphs[idx]
            if nxt_p.style.name.startswith('Heading') or nxt_p.text.strip().startswith('B.') or "BAB B" in nxt_p.text:
                break
            delete_paragraph(nxt_p)
            idx += 1
            
        new_p = insert_paragraph_after(p_rtm, "Tabel Requirements Traceability Matrix (RTM):", style="Normal")
        
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        headers = ["ID Requirement", "Deskripsi Requirement", "Test Case ID"]
        for i, header in enumerate(headers):
            table.cell(0, i).text = header
            set_cell_background(table.cell(0, i), '0E748C')
            run = table.cell(0, i).paragraphs[0].runs[0]
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)

        rtm_data = [
            ("FR1", "Pengguna dapat melihat halaman utama (Landing Page)", "BB-001 s/d BB-010"),
            ("FR2", "Pengguna dapat melakukan login dan logout", "BB-011 s/d BB-025, WB-010 s/d WB-014, WB-022 s/d WB-029"),
            ("FR3", "Pengguna dapat mendaftar (Registrasi)", "BB-026 s/d BB-040, WB-018 s/d WB-021, WB-001 s/d WB-005"),
            ("FR4", "Pengguna dapat mengisi assessment & rekomendasi", "BB-051 s/d BB-070, WB-015 s/d WB-017"),
            ("FR5", "Pengguna dapat menggunakan fitur Chat", "BB-071 s/d BB-080"),
            ("FR6", "Pengguna dapat melakukan Booking Jadwal", "BB-081 s/d BB-090"),
            ("FR7", "Konselor dapat mengelola ketersediaan jadwal", "BB-091 s/d BB-095, WB-030 s/d WB-031"),
            ("NFR1", "Tampilan aplikasi responsif", "BB-002, BB-003, BB-004"),
            ("NFR2", "Keamanan token & Role-based Access", "BB-096 s/d BB-100, WB-006 s/d WB-009"),
        ]
        
        for row_data in rtm_data:
            row_cells = table.add_row().cells
            for i, val in enumerate(row_data):
                row_cells[i].text = val
                
        tbl_element = table._element
        doc._body._element.remove(tbl_element)
        new_p._p.addnext(tbl_element)

    if nfr_heading_idx != -1:
        p_nfr = doc.paragraphs[nfr_heading_idx]
        idx = nfr_heading_idx + 1
        while idx < len(doc.paragraphs):
            nxt_p = doc.paragraphs[idx]
            if nxt_p.style.name.startswith('Heading') or "4." in nxt_p.text or nxt_p.text.strip().startswith('4 '):
                break
            delete_paragraph(nxt_p)
            idx += 1
            
        nfr_items = [
            "NFR1: Tampilan aplikasi responsif (Mobile View & Desktop).",
            "NFR2: Keamanan token JWT dengan pengecekan expiry otomatis.",
            "NFR3: Role-based access control (RBAC) ketat antara Pelajar, Konselor, dan Admin.",
            "NFR4: Validasi input form secara realtime di client-side.",
            "NFR5: Umpan balik antarmuka yang informatif (Loading state, error messages, toast notification)."
        ]
        curr_p = p_nfr
        for item in nfr_items:
            curr_p = insert_paragraph_after(curr_p, item, style="List Paragraph")
            
    if fr_heading_idx != -1:
        p_fr = doc.paragraphs[fr_heading_idx]
        idx = fr_heading_idx + 1
        while idx < len(doc.paragraphs):
            nxt_p = doc.paragraphs[idx]
            if nxt_p.style.name.startswith('Heading') or "3." in nxt_p.text or nxt_p.text.strip().startswith('3 '):
                break
            delete_paragraph(nxt_p)
            idx += 1
            
        fr_items = [
            "FR1: Pengguna dapat mengakses halaman utama dan navigasi global.",
            "FR2: Pengguna dapat masuk (login) ke dalam sistem dan keluar (logout) dengan aman.",
            "FR3: Pengguna baru dapat mendaftar (registrasi) sebagai Pelajar atau Konselor dengan foto profil.",
            "FR4: Pelajar dapat mengisi kuesioner assessment emosi dan menerima rekomendasi.",
            "FR5: Pelajar dan Konselor dapat berkomunikasi real-time melalui fitur Chat Room.",
            "FR6: Pelajar dapat memesan sesi konseling (Booking) dengan konselor yang tersedia.",
            "FR7: Konselor dapat mengelola jadwal ketersediaan untuk sesi konseling."
        ]
        curr_p = p_fr
        for item in fr_items:
            curr_p = insert_paragraph_after(curr_p, item, style="List Paragraph")

# ======================= TAHAP 2: BAB B =======================
def update_bab_b(doc):
    target_p = None
    bab_c_p = None
    for p in doc.paragraphs:
        text = p.text.strip().lower()
        if "fitur yang diuji" in text and ("2." in text or "2 " in text):
            target_p = p
        elif text and ("bab c" in text or "c." in text.split()[0]):
            if bab_c_p is None:
                bab_c_p = p
            
    if target_p and bab_c_p:
        parent = target_p._element.getparent()
        start_idx = parent.index(target_p._element)
        end_idx = parent.index(bab_c_p._element)
        
        # Hapus elemen lama di antara "Fitur yang diuji" dan "BAB C"
        for el in parent[start_idx+1:end_idx]:
            parent.remove(el)
            
        features = [
            "1. Landing Page (Global Navbar, Hero Section, Responsiveness)",
            "2. Authentication (Login, Registrasi Multi-step, Role-based redirect, JWT Management, Image Upload)",
            "3. Home & Layout (Protected routes, Sidebar navigasi, Token expiry guard)",
            "4. Assessment & Recommendation (Kuesioner emosi, Perhitungan intensitas, Hasil rekomendasi artikel/video)",
            "5. Chat System (Real-time socket.io, Pencarian konselor, Chat room)",
            "6. Booking System (Pemesanan sesi, Approval/Rejection oleh Konselor, Riwayat booking)",
            "7. Counselor Schedule Management (CRUD ketersediaan jadwal)"
        ]
        
        curr_el = target_p._element
        for feat in features:
            new_p = OxmlElement("w:p")
            curr_el.addnext(new_p)
            new_para = docx.text.paragraph.Paragraph(new_p, target_p._parent)
            new_para.add_run(feat)
            new_para.style = "List Paragraph"
            curr_el = new_p

# ======================= TAHAP 3: BAB C =======================
def add_test_case_table(doc, headers, data, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    
    # Header
    hdr_cells = table.rows[0].cells
    for i, title_text in enumerate(headers):
        hdr_cells[i].text = title_text
        if col_widths and i < len(col_widths):
            hdr_cells[i].width = col_widths[i]
        set_cell_background(hdr_cells[i], '0E748C') # sky-700
        set_cell_margins(hdr_cells[i], top=100, bottom=100, left=100, right=100)
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.name = 'Arial'
        run.font.size = Pt(9.0)
        
    # Rows
    for row_idx, row_data in enumerate(data):
        row_cells = table.add_row().cells
        bg_color = 'F0FDFA' if row_idx % 2 == 0 else 'FFFFFF'
        
        for col_idx, cell_value in enumerate(row_data):
            row_cells[col_idx].text = str(cell_value)
            if col_widths and col_idx < len(col_widths):
                row_cells[col_idx].width = col_widths[col_idx]
            set_cell_background(row_cells[col_idx], bg_color)
            set_cell_margins(row_cells[col_idx], top=80, bottom=80, left=80, right=80)
            
            p = row_cells[col_idx].paragraphs[0]
            if col_idx == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
            if len(p.runs) > 0:
                run = p.runs[0]
                run.font.name = 'Arial'
                run.font.size = Pt(8.5)
                if col_idx == 0:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(75, 85, 99)
            row_cells[col_idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

def update_bab_c(doc):
    target_p = None
    for p in doc.paragraphs:
        text = p.text.strip().lower()
        if "test case development" in text or "skenario pengujian" in text:
            target_p = p
            
    if not target_p:
        # Fallback to the last paragraph that looks like Bab C heading
        for p in doc.paragraphs:
            if "bab c" in p.text.strip().lower() or "c." in p.text.strip().lower():
                target_p = p
                
    if not target_p:
        target_p = doc.paragraphs[-1] # Extreme fallback
        
    # Clear all elements after target_p except sectPr
    parent = target_p._element.getparent()
    idx = parent.index(target_p._element)
    for el in parent[idx+1:]:
        if el.tag.endswith('sectPr'):
            continue
        parent.remove(el)
        
    # Tambahkan Sub Heading untuk Black Box
    p_bb = doc.add_paragraph()
    run_bb = p_bb.add_run("1. Skenario Uji Black Box (100 Test Cases)")
    run_bb.font.bold = True
    run_bb.font.size = Pt(12)
    
    # Tambahkan Tabel Black Box
    headers_bb = ["Test ID", "Deskripsi Pengujian", "Langkah Pengujian", "Data Input", "Expected Output"]
    widths_bb = [Inches(0.6), Inches(1.5), Inches(1.8), Inches(1.2), Inches(1.8)]
    add_test_case_table(doc, headers_bb, bb_cases_raw, widths_bb)
    
    doc.add_paragraph() # Spacer
    
    # Tambahkan Sub Heading untuk White Box
    p_wb = doc.add_paragraph()
    run_wb = p_wb.add_run("2. Skenario Uji White Box (31 Test Cases)")
    run_wb.font.bold = True
    run_wb.font.size = Pt(12)
    
    # Tambahkan Tabel White Box
    headers_wb = ["Test ID", "Deskripsi Pengujian", "Potongan Kode yang Diuji", "Potongan Kode Uji"]
    widths_wb = [Inches(0.6), Inches(2.0), Inches(2.0), Inches(2.0)]
    add_test_case_table(doc, headers_wb, wb_cases, widths_wb)


# ======================= MAIN =======================
def process_full_document(input_path, output_path):
    print(f"Membaca dokumen {input_path}...")
    try:
        doc = Document(input_path)
    except Exception as e:
        print(f"Gagal membaca dokumen: {e}")
        sys.exit(1)
        
    print("Memproses Tahap 1: Bab A (FR, NFR, RTM)...")
    update_bab_a(doc)
    
    print("Memproses Tahap 2: Bab B (Fitur yang diuji)...")
    update_bab_b(doc)
    
    print("Memproses Tahap 3: Bab C (Test Case Development Table BB 100 cases & WB 31 cases)...")
    update_bab_c(doc)
    
    try:
        # Save requires orientation handling to avoid landscape issues if any, but python-docx keeps original orientation by default
        doc.save(output_path)
        print(f"Selesai! Dokumen berhasil disimpan sebagai {output_path}")
    except Exception as e:
        print(f"Gagal menyimpan dokumen: {e}")

if __name__ == "__main__":
    input_file = "STLC Pada Website Konseling Sejiwa App.docx"
    output_file = "STLC_Final.docx"
    process_full_document(input_file, output_file)
